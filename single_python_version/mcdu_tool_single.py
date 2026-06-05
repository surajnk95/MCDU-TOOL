from __future__ import annotations

import csv
import io
import json
import math
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
from pathlib import Path
from tkinter import BOTH, BOTTOM, END, HORIZONTAL, LEFT, RIGHT, TOP, X, Y, Canvas, DoubleVar, Entry, Frame, Label, StringVar, Tk
from tkinter import filedialog, messagebox, ttk
from typing import Any

import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from PIL import Image, ImageEnhance, ImageFilter, ImageOps, ImageTk


ROWS = 13
COLS = 40
FIRST_DATA_COL = 1
LAST_DATA_COL = 38
SCREEN_W = 1600
MIN_SCREEN_H = 900
MAX_SCREEN_H = 1400

DEFAULT_CORRECTIONS = {
    "-000°/0.0NN": " 000°/0.0NM                 000°/0.0NM  ",
    "000°/0.0NH 000°/0.0NM": " 000°/0.0NM                 000°/0.0NM  ",
    "000°/0.0NH 000°/0..0N": " 000°/0.0NM                 000°/0.0NM  ",
    "000°/0.0NH 000°/0.0": " 000°/0.0NM                 000°/0.0NM  ",
    "000°/0.0NH 000°/0.01": " 000°/0.0NM                 000°/0.0NM  ",
    "000°/0.0NH 000°/0.ONM": " 000°/0.0NM                 000°/0.0NM  ",
    "000°/0.0NM 000°/0.00": " 000°/0.0NM                 000°/0.0NM  ",
    "000°70.ONM -000°/0.0NN": " 000°/0.0NM                 000°/0.0NM  ",
    "000°/0.ONM 000°/0.ON": " 000°/0.0NM                 000°/0.0NM  ",
    "000°/0.ONN 000°/0..OM": " 000°/0.0NM                 000°/0.0NM  ",
    "000°70.0NH 000°/0.0MH": " 000°/0.0NM                 000°/0.0NM  ",
    "000°70.0NN -000°/0.0NN": " 000°/0.0NM                 000°/0.0NM  ",
    "6PSL 6PSR": "  GPS L                         GPS R   ",
    "6PSL PSR": "  GPS L                         GPS R   ",
    "<INDEX OFF->": " <INDEX                       OFF<>ON>  ",
    "<INDEX OFF>": " <INDEX                       OFF<>ON>  ",
    "CPS-ANRUL CPS-ANRYR": "  GPS-ANRU L               GPS-ANRU R   ",
    "GG SENSORSELECT": " ------------------------ SENSOR SELECT ",
    "GS SENSORSELECT": " ------------------------ SENSOR SELECT ",
    "OPS-TRUL OPS-3RUR": "  GPS-IRU L                  GPS-IRU R  ",
    "OPS-TRUL OPS-IRUB": "  GPS-IRU L                  GPS-IRU R  ",
    "OC SSEECT": " ------------------------ SENSOR SELECT ",
    "POSREF 4": "                POS REF           4/4   ",
    "POSREF MM": "                POS REF           4/4   ",
    "POSREFMM": "                POS REF           4/4   ",
    "POSREF Y": "                POS REF           4/4   ",
    "R010": "  RADIO                                 ",
    "RUL RUR": "  IRU L                         IRU R   ",
    "RUL IRUR": "  IRU L                         IRU R   ",
    "SPS-ANRUL GPS-ANRYR": "  GPS-ANRU L               GPS-ANRU R   ",
    "SPS-ANRUL GPS-ANRUR": "  GPS-ANRU L               GPS-ANRU R   ",
}


def app_dir() -> Path:
    if getattr(sys, "frozen", False):
        return Path(sys.executable).resolve().parent
    return Path(__file__).resolve().parent


APP_DIR = app_dir()
CORRECTIONS = APP_DIR / "corrections.json"
TEMPLATES = APP_DIR / "templates.json"
EXPORTS = APP_DIR / "exports"


def tesseract_cmd() -> str:
    configured = os.environ.get("TESSERACT_CMD")
    if configured:
        return configured
    found = shutil.which("tesseract")
    if found:
        return found
    homebrew = Path("/opt/homebrew/bin/tesseract")
    if homebrew.exists():
        return str(homebrew)
    return "tesseract"


TESSERACT = tesseract_cmd()


def ensure_dirs() -> None:
    EXPORTS.mkdir(parents=True, exist_ok=True)
    if not CORRECTIONS.exists():
        CORRECTIONS.write_text(json.dumps(DEFAULT_CORRECTIONS, indent=2, sort_keys=True), encoding="utf-8")
    if not TEMPLATES.exists():
        TEMPLATES.write_text("{}", encoding="utf-8")


def edge_length(a: dict[str, float], b: dict[str, float]) -> float:
    return math.hypot(a["x"] - b["x"], a["y"] - b["y"])


def lerp(a: float, b: float, t: float) -> float:
    return a + (b - a) * t


def lerp_point(a: dict[str, float], b: dict[str, float], t: float) -> dict[str, float]:
    return {"x": lerp(a["x"], b["x"], t), "y": lerp(a["y"], b["y"], t)}


def bilinear(corners: list[dict[str, float]], u: float, v: float) -> dict[str, float]:
    top = lerp_point(corners[0], corners[1], u)
    bottom = lerp_point(corners[3], corners[2], u)
    return lerp_point(top, bottom, v)


def screen_size_from_corners(corners: list[dict[str, float]]) -> tuple[int, int]:
    top = edge_length(corners[0], corners[1])
    right = edge_length(corners[1], corners[2])
    bottom = edge_length(corners[2], corners[3])
    left = edge_length(corners[3], corners[0])
    width = max(1.0, (top + bottom) * 0.5)
    height = max(1.0, (left + right) * 0.5)
    aspect = width / height
    normalized_height = int(round(SCREEN_W / aspect))
    return SCREEN_W, max(MIN_SCREEN_H, min(MAX_SCREEN_H, normalized_height))


def solve_linear_system(matrix: list[list[float]], values: list[float]) -> list[float]:
    n = len(values)
    augmented = [row[:] + [values[i]] for i, row in enumerate(matrix)]
    for pivot in range(n):
        max_row = max(range(pivot, n), key=lambda row: abs(augmented[row][pivot]))
        augmented[pivot], augmented[max_row] = augmented[max_row], augmented[pivot]
        pivot_value = augmented[pivot][pivot]
        if abs(pivot_value) < 1e-12:
            raise ValueError("Screen corners are too close together.")
        for col in range(pivot, n + 1):
            augmented[pivot][col] /= pivot_value
        for row in range(n):
            if row == pivot:
                continue
            factor = augmented[row][pivot]
            for col in range(pivot, n + 1):
                augmented[row][col] -= factor * augmented[pivot][col]
    return [augmented[row][n] for row in range(n)]


def perspective_coefficients(src: list[dict[str, float]], dst_size: tuple[int, int]) -> list[float]:
    width, height = dst_size
    dst = [
        {"x": 0.0, "y": 0.0},
        {"x": float(width), "y": 0.0},
        {"x": float(width), "y": float(height)},
        {"x": 0.0, "y": float(height)},
    ]
    matrix: list[list[float]] = []
    values: list[float] = []
    for d, s in zip(dst, src):
        x = d["x"]
        y = d["y"]
        u = s["x"]
        v = s["y"]
        matrix.append([x, y, 1, 0, 0, 0, -u * x, -u * y])
        values.append(u)
        matrix.append([0, 0, 0, x, y, 1, -v * x, -v * y])
        values.append(v)
    return solve_linear_system(matrix, values)


def warp_screen(image: Image.Image, corners: list[dict[str, float]]) -> Image.Image:
    screen_size = screen_size_from_corners(corners)
    coeffs = perspective_coefficients(corners, screen_size)
    return image.transform(screen_size, Image.Transform.PERSPECTIVE, coeffs, Image.Resampling.BICUBIC)


def preprocess_for_ocr(image: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(image)
    gray = ImageOps.invert(gray)
    gray = ImageEnhance.Contrast(gray).enhance(2.8)
    return gray.filter(ImageFilter.SHARPEN)


def clean_ocr_text(text: str) -> str:
    text = text.replace("|", "I")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def empty_grid() -> list[list[str]]:
    return [["" for _ in range(COLS)] for _ in range(ROWS)]


def clamp_data_col(col: int) -> int:
    return max(FIRST_DATA_COL, min(LAST_DATA_COL, col))


def nearest_data_col(x: float, cell_w: float) -> int:
    return clamp_data_col(int(round(x / cell_w)))


def normalize_grid_guards(grid: list[list[str]]) -> list[list[str]]:
    normalized = [[cell for cell in row[:COLS]] + [""] * max(0, COLS - len(row)) for row in grid[:ROWS]]
    while len(normalized) < ROWS:
        normalized.append(["" for _ in range(COLS)])
    for row in normalized:
        row[0] = ""
        row[COLS - 1] = ""
    return normalized


def connected_components(mask: np.ndarray) -> list[dict[str, Any]]:
    height, width = mask.shape
    seen = np.zeros(mask.shape, dtype=bool)
    components: list[dict[str, Any]] = []
    for start_y in range(height):
        for start_x in range(width):
            if seen[start_y, start_x] or not mask[start_y, start_x]:
                continue
            stack = [(start_x, start_y)]
            seen[start_y, start_x] = True
            xs: list[int] = []
            ys: list[int] = []
            while stack:
                x, y = stack.pop()
                xs.append(x)
                ys.append(y)
                for next_x, next_y in ((x + 1, y), (x - 1, y), (x, y + 1), (x, y - 1)):
                    if (
                        0 <= next_x < width
                        and 0 <= next_y < height
                        and not seen[next_y, next_x]
                        and mask[next_y, next_x]
                    ):
                        seen[next_y, next_x] = True
                        stack.append((next_x, next_y))
            count = len(xs)
            if count < 250:
                continue
            x1 = min(xs)
            x2 = max(xs) + 1
            y1 = min(ys)
            y2 = max(ys) + 1
            components.append(
                {
                    "count": count,
                    "x1": x1,
                    "y1": y1,
                    "x2": x2,
                    "y2": y2,
                    "xs": xs,
                    "ys": ys,
                    "fill": count / max(1, (x2 - x1) * (y2 - y1)),
                }
            )
    return components


def component_corners(component: dict[str, Any], scale: float) -> list[dict[str, float]]:
    xs = np.asarray(component.get("xs", []), dtype=np.float32)
    ys = np.asarray(component.get("ys", []), dtype=np.float32)
    if xs.size < 50 or ys.size < 50:
        raise ValueError("Not enough points to refine screen corners.")

    def fit_x_edge(edge: str) -> tuple[float, float]:
        samples_y: list[float] = []
        samples_x: list[float] = []
        bins = np.linspace(float(component["y1"]), float(component["y2"]), 36)
        for low, high in zip(bins[:-1], bins[1:]):
            mask = (ys >= low) & (ys < high)
            values = xs[mask]
            if values.size < 8:
                continue
            samples_y.append((low + high) * 0.5)
            samples_x.append(float(np.percentile(values, 1 if edge == "left" else 99)))
        if len(samples_y) < 6:
            raise ValueError("Could not fit vertical display edge.")
        slope, intercept = np.polyfit(np.asarray(samples_y), np.asarray(samples_x), 1)
        return float(slope), float(intercept)

    def fit_y_edge(edge: str) -> tuple[float, float]:
        samples_x: list[float] = []
        samples_y: list[float] = []
        bins = np.linspace(float(component["x1"]), float(component["x2"]), 36)
        for low, high in zip(bins[:-1], bins[1:]):
            mask = (xs >= low) & (xs < high)
            values = ys[mask]
            if values.size < 8:
                continue
            samples_x.append((low + high) * 0.5)
            samples_y.append(float(np.percentile(values, 1 if edge == "top" else 99)))
        if len(samples_x) < 6:
            raise ValueError("Could not fit horizontal display edge.")
        slope, intercept = np.polyfit(np.asarray(samples_x), np.asarray(samples_y), 1)
        return float(slope), float(intercept)

    left = fit_x_edge("left")
    right = fit_x_edge("right")
    top = fit_y_edge("top")
    bottom = fit_y_edge("bottom")

    def intersect_xedge_yedge(x_edge: tuple[float, float], y_edge: tuple[float, float]) -> dict[str, float]:
        # x = a*y + b, y = c*x + d
        a, b = x_edge
        c, d = y_edge
        denominator = 1.0 - c * a
        if abs(denominator) < 1e-6:
            raise ValueError("Detected display edges are nearly parallel.")
        y = (c * b + d) / denominator
        x = a * y + b
        return {"x": float(x), "y": float(y)}

    corners = [
        intersect_xedge_yedge(left, top),
        intersect_xedge_yedge(right, top),
        intersect_xedge_yedge(right, bottom),
        intersect_xedge_yedge(left, bottom),
    ]

    width_top = edge_length(corners[0], corners[1])
    width_bottom = edge_length(corners[3], corners[2])
    height_left = edge_length(corners[0], corners[3])
    height_right = edge_length(corners[1], corners[2])
    if min(width_top, width_bottom, height_left, height_right) < 25:
        raise ValueError("Refined screen corners are too close together.")

    pad_x = max(0.5, (component["x2"] - component["x1"]) * 0.001)
    pad_y = max(0.5, (component["y2"] - component["y1"]) * 0.001)
    center_x = float(np.mean(xs))
    center_y = float(np.mean(ys))
    refined: list[dict[str, float]] = []
    for corner in corners:
        x = corner["x"] + pad_x if corner["x"] < center_x else corner["x"] - pad_x
        y = corner["y"] + pad_y if corner["y"] < center_y else corner["y"] - pad_y
        refined.append({"x": x / scale, "y": y / scale})
    return refined


def detect_display(image: Image.Image) -> list[dict[str, float]]:
    scale = min(1.0, 900 / max(image.size))
    small = image.resize((max(1, round(image.width * scale)), max(1, round(image.height * scale))))
    arr = np.asarray(small).astype(np.float32)
    gray = arr[:, :, 0] * 0.299 + arr[:, :, 1] * 0.587 + arr[:, :, 2] * 0.114
    components = connected_components(gray < 58)
    if not components:
        raise ValueError("Could not find a dark MCMDU display region.")

    scored: list[tuple[float, dict[str, Any]]] = []
    image_area = small.width * small.height
    for component in components:
        width = component["x2"] - component["x1"]
        height = component["y2"] - component["y1"]
        area = width * height
        aspect = width / max(1, height)
        touches_edge = (
            component["x1"] <= 2
            or component["y1"] <= 2
            or component["x2"] >= small.width - 2
            or component["y2"] >= small.height - 2
        )
        if area < image_area * 0.08 or aspect < 0.75 or aspect > 2.2:
            continue
        score = component["count"] * min(1.0, component["fill"] * 1.4)
        if touches_edge:
            score *= 0.35
        scored.append((score, component))
    if not scored:
        raise ValueError("Could not isolate the black display. Drag the four corners manually.")

    _, best = max(scored, key=lambda item: item[0])
    pad_x = max(1, round((best["x2"] - best["x1"]) * 0.003))
    pad_y = max(1, round((best["y2"] - best["y1"]) * 0.003))
    x1 = max(0, best["x1"] + pad_x)
    y1 = max(0, best["y1"] + pad_y)
    x2 = min(small.width, best["x2"] - pad_x)
    y2 = min(small.height, best["y2"] - pad_y)
    inv = 1 / scale
    return [
        {"x": x1 * inv, "y": y1 * inv},
        {"x": x2 * inv, "y": y1 * inv},
        {"x": x2 * inv, "y": y2 * inv},
        {"x": x1 * inv, "y": y2 * inv},
    ]


def run_tesseract_tsv(image: Image.Image) -> list[dict[str, Any]]:
    with tempfile.TemporaryDirectory() as temp_dir:
        source = Path(temp_dir) / "screen.png"
        image.save(source)
        command = [
            TESSERACT,
            str(source),
            "stdout",
            "--psm",
            "11",
            "-l",
            "eng",
            "-c",
            "tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789°/.-<>",
            "tsv",
        ]
        completed = subprocess.run(command, capture_output=True, text=True, check=False)
        if completed.returncode != 0:
            raise RuntimeError(completed.stderr.strip() or "Tesseract OCR failed.")
        rows = csv.DictReader(io.StringIO(completed.stdout), delimiter="\t")
        words: list[dict[str, Any]] = []
        for row in rows:
            text = (row.get("text") or "").strip()
            if not text:
                continue
            try:
                conf = float(row.get("conf") or -1)
            except ValueError:
                conf = -1
            if conf < 0:
                continue
            words.append(
                {
                    "text": text,
                    "conf": conf,
                    "left": int(row.get("left") or 0),
                    "top": int(row.get("top") or 0),
                    "width": int(row.get("width") or 0),
                    "height": int(row.get("height") or 0),
                }
            )
        return words


def run_tesseract_boxes(image: Image.Image) -> list[dict[str, Any]]:
    width, height = image.size
    with tempfile.TemporaryDirectory() as temp_dir:
        source = Path(temp_dir) / "screen.png"
        image.save(source)
        command = [
            TESSERACT,
            str(source),
            "stdout",
            "--psm",
            "6",
            "-l",
            "eng",
            "-c",
            "tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789°/.-<>",
            "makebox",
        ]
        completed = subprocess.run(command, capture_output=True, text=True, check=False)
        if completed.returncode != 0:
            return []
        boxes: list[dict[str, Any]] = []
        for line in completed.stdout.splitlines():
            parts = line.split()
            if len(parts) < 5:
                continue
            char = clean_ocr_text(parts[0])
            if not char:
                continue
            try:
                left = int(parts[1])
                bottom = int(parts[2])
                right = int(parts[3])
                top = int(parts[4])
            except ValueError:
                continue
            boxes.append(
                {
                    "text": char[:1],
                    "left": left,
                    "top": max(0, height - top),
                    "width": max(0, right - left),
                    "height": max(0, top - bottom),
                }
            )
        return boxes


def place_char(grid: list[list[str]], char_box: dict[str, Any], screen_size: tuple[int, int]) -> bool:
    text = clean_ocr_text(str(char_box["text"]))[:1]
    if not text:
        return False
    screen_w, screen_h = screen_size
    cell_w = screen_w / COLS
    cell_h = screen_h / ROWS
    box_w = float(char_box.get("width") or 0)
    box_h = float(char_box.get("height") or 0)
    if box_w > cell_w * 2.15 or box_h > cell_h * 1.35 or box_w < 2 or box_h < 2:
        return False
    center_x = float(char_box["left"]) + box_w * 0.5
    center_y = float(char_box["top"]) + box_h * 0.5
    row = max(0, min(ROWS - 1, int(center_y / cell_h)))
    col = clamp_data_col(int(center_x / cell_w))
    if grid[row][col] and grid[row][col] != text:
        for offset in (-1, 1, -2, 2):
            next_col = col + offset
            if FIRST_DATA_COL <= next_col <= LAST_DATA_COL and not grid[row][next_col]:
                col = next_col
                break
    if not grid[row][col]:
        grid[row][col] = text
        return True
    return False


def place_word(grid: list[list[str]], word: dict[str, Any], screen_size: tuple[int, int]) -> None:
    text = clean_ocr_text(str(word["text"]))
    if not text:
        return
    screen_w, screen_h = screen_size
    cell_w = screen_w / COLS
    cell_h = screen_h / ROWS
    row = max(0, min(ROWS - 1, int((word["top"] + word["height"] * 0.5) / cell_h)))
    start_col = nearest_data_col(float(word["left"]), cell_w)
    compact = text.replace(" ", "")
    while start_col <= LAST_DATA_COL and grid[row][start_col]:
        start_col += 1
    if start_col > LAST_DATA_COL:
        return
    for index, char in enumerate(compact[: LAST_DATA_COL - start_col + 1]):
        col = start_col + index
        if not grid[row][col]:
            grid[row][col] = char


def load_corrections() -> dict[str, str]:
    ensure_dirs()
    try:
        data = json.loads(CORRECTIONS.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        data = {}
    return {str(key): str(value) for key, value in data.items()}


def save_corrections(corrections: dict[str, str]) -> None:
    ensure_dirs()
    CORRECTIONS.write_text(json.dumps(corrections, indent=2, sort_keys=True), encoding="utf-8")


def apply_corrections(grid: list[list[str]]) -> list[list[str]]:
    corrections = load_corrections()
    updated = [[cell for cell in row] for row in grid]
    for row_index, row in enumerate(updated):
        row_text = "".join(cell or " " for cell in row)
        normalized = re.sub(r"\s+", " ", row_text).strip()
        if normalized in corrections:
            corrected = corrections[normalized][:COLS].ljust(COLS)
            updated[row_index] = list(corrected)
    return normalize_grid_guards(updated)


def grid_character_count(grid: list[list[str]]) -> int:
    return sum(1 for row in grid for cell in row[FIRST_DATA_COL : LAST_DATA_COL + 1] if cell)


def analyze_image(image: Image.Image, corners: list[dict[str, float]]) -> tuple[list[list[str]], Image.Image, int, int]:
    warped = warp_screen(image, corners)
    processed = preprocess_for_ocr(warped)
    screen_size = warped.size
    boxes = run_tesseract_boxes(processed)
    words = run_tesseract_tsv(processed)
    grid = empty_grid()
    for word in words:
        place_word(grid, word, screen_size)
    if grid_character_count(grid) < 8:
        for box in boxes:
            place_char(grid, box, screen_size)
    return normalize_grid_guards(apply_corrections(grid)), warped, len(boxes), len(words)


def add_table_header(table: Any) -> None:
    header = table.rows[0].cells
    header[0].text = "Row"
    for col in range(COLS):
        header[col + 1].text = str(col) if FIRST_DATA_COL <= col <= LAST_DATA_COL else ""
        for paragraph in header[col + 1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_docx(grid: list[list[str]], output_path: Path) -> None:
    document = Document()
    document.add_heading("777-9 MCMDU Grid Extraction", level=1)
    document.add_paragraph("Rows are numbered 1 to 13. Screen columns 2 to 39 are labelled 1 to 38.")
    table = document.add_table(rows=ROWS + 1, cols=COLS + 1)
    table.style = "Table Grid"
    add_table_header(table)
    normalized = normalize_grid_guards(grid)
    for row_index in range(ROWS):
        cells = table.rows[row_index + 1].cells
        cells[0].text = str(row_index + 1)
        for col_index in range(COLS):
            value = "" if col_index in (0, COLS - 1) else str(normalized[row_index][col_index])
            cells[col_index + 1].text = value.strip()
            cells[col_index + 1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[col_index + 1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(7)
    document.save(output_path)


class McmduDesktopTool:
    def __init__(self, root: Tk) -> None:
        self.root = root
        self.root.title("777-9 MCMDU Grid Tool - Single Python")
        self.image: Image.Image | None = None
        self.warped: Image.Image | None = None
        self.view_image: Image.Image | None = None
        self.photo_image: ImageTk.PhotoImage | None = None
        self.image_path: Path | None = None
        self.corners: list[dict[str, float]] = []
        self.display_mode = "photo"
        self.drag_corner: int | None = None
        self.scale = 1.0
        self.offset_x = 0
        self.offset_y = 0
        self.inset_vars = {
            "left": DoubleVar(value=0.0),
            "right": DoubleVar(value=0.0),
            "top": DoubleVar(value=0.0),
            "bottom": DoubleVar(value=0.0),
        }
        self.grid = empty_grid()
        self.source_grid = empty_grid()
        self.entries: list[list[Entry]] = []
        self.status_var = StringVar(value=f"Ready. Training data folder: {APP_DIR}")
        self.build_ui()

    def build_ui(self) -> None:
        toolbar = Frame(self.root)
        toolbar.pack(side=TOP, fill=X, padx=8, pady=6)
        ttk.Button(toolbar, text="Choose Image", command=self.choose_image).pack(side=LEFT, padx=3)
        ttk.Button(toolbar, text="Auto Detect Display", command=self.auto_detect).pack(side=LEFT, padx=3)
        ttk.Button(toolbar, text="Photo View", command=lambda: self.set_mode("photo")).pack(side=LEFT, padx=3)
        ttk.Button(toolbar, text="Flattened View", command=lambda: self.set_mode("flat")).pack(side=LEFT, padx=3)
        ttk.Button(toolbar, text="Analyze Grid", command=self.analyze).pack(side=LEFT, padx=3)
        ttk.Button(toolbar, text="Remember Corrections", command=self.remember_corrections).pack(side=LEFT, padx=3)
        ttk.Button(toolbar, text="Export Word File", command=self.export_word).pack(side=LEFT, padx=3)

        inset_bar = Frame(self.root)
        inset_bar.pack(side=TOP, fill=X, padx=8, pady=(0, 6))
        for label, key in (("Left inset", "left"), ("Right inset", "right"), ("Top inset", "top"), ("Bottom inset", "bottom")):
            Label(inset_bar, text=label).pack(side=LEFT, padx=(8, 3))
            ttk.Scale(
                inset_bar,
                from_=0,
                to=12,
                orient=HORIZONTAL,
                variable=self.inset_vars[key],
                command=lambda _value: self.render_canvas(),
                length=120,
            ).pack(side=LEFT, padx=(0, 6))

        content = Frame(self.root)
        content.pack(side=TOP, fill=BOTH, expand=True, padx=8, pady=4)

        left = Frame(content)
        left.pack(side=LEFT, fill=BOTH, expand=True)
        self.canvas = Canvas(left, bg="#101415", height=520)
        self.canvas.pack(side=TOP, fill=BOTH, expand=True)
        self.canvas.bind("<ButtonPress-1>", self.on_press)
        self.canvas.bind("<B1-Motion>", self.on_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_release)
        self.canvas.bind("<Configure>", lambda _event: self.render_canvas())

        right = Frame(content)
        right.pack(side=RIGHT, fill=BOTH, padx=(8, 0))
        Label(right, text="Extracted Grid").pack(anchor="w")

        grid_wrap = Frame(right)
        grid_wrap.pack(side=TOP, fill=BOTH, expand=True)
        self.grid_canvas = Canvas(grid_wrap, width=780, height=460)
        self.grid_canvas.pack(side=LEFT, fill=BOTH, expand=True)
        y_scroll = ttk.Scrollbar(grid_wrap, orient="vertical", command=self.grid_canvas.yview)
        y_scroll.pack(side=RIGHT, fill=Y)
        x_scroll = ttk.Scrollbar(right, orient=HORIZONTAL, command=self.grid_canvas.xview)
        x_scroll.pack(side=BOTTOM, fill=X)
        self.grid_canvas.configure(yscrollcommand=y_scroll.set, xscrollcommand=x_scroll.set)
        self.grid_frame = Frame(self.grid_canvas)
        self.grid_canvas.create_window((0, 0), window=self.grid_frame, anchor="nw")
        self.grid_frame.bind(
            "<Configure>",
            lambda _event: self.grid_canvas.configure(scrollregion=self.grid_canvas.bbox("all")),
        )
        self.build_grid()

        status = ttk.Label(self.root, textvariable=self.status_var, anchor="w")
        status.pack(side=BOTTOM, fill=X, padx=8, pady=4)

    def build_grid(self) -> None:
        Label(self.grid_frame, text="Row", width=4).grid(row=0, column=0, sticky="nsew")
        for col in range(COLS):
            label = str(col) if FIRST_DATA_COL <= col <= LAST_DATA_COL else ""
            Label(self.grid_frame, text=label, width=2).grid(row=0, column=col + 1, sticky="nsew")
        self.entries = []
        for row in range(ROWS):
            Label(self.grid_frame, text=str(row + 1), width=4).grid(row=row + 1, column=0, sticky="nsew")
            entry_row: list[Entry] = []
            for col in range(COLS):
                entry = Entry(self.grid_frame, width=2, justify="center", font=("Courier", 10))
                entry.grid(row=row + 1, column=col + 1, sticky="nsew")
                if col < FIRST_DATA_COL or col > LAST_DATA_COL:
                    entry.configure(state="disabled", disabledbackground="#eef2f3")
                entry_row.append(entry)
            self.entries.append(entry_row)

    def set_grid(self, grid: list[list[str]]) -> None:
        self.grid = normalize_grid_guards(grid)
        for row in range(ROWS):
            for col in range(COLS):
                entry = self.entries[row][col]
                if col < FIRST_DATA_COL or col > LAST_DATA_COL:
                    continue
                entry.delete(0, END)
                if self.grid[row][col]:
                    entry.insert(0, self.grid[row][col][:1])

    def get_grid(self) -> list[list[str]]:
        grid = empty_grid()
        for row in range(ROWS):
            for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
                grid[row][col] = self.entries[row][col].get().strip()[:1]
        return normalize_grid_guards(grid)

    def choose_image(self) -> None:
        filename = filedialog.askopenfilename(
            title="Choose MCMDU Image",
            filetypes=[("Images", "*.jpg *.jpeg *.png *.bmp *.tif *.tiff"), ("All files", "*.*")],
        )
        if not filename:
            return
        self.image_path = Path(filename)
        self.image = ImageOps.exif_transpose(Image.open(filename)).convert("RGB")
        w, h = self.image.size
        margin_x = w * 0.18
        margin_y = h * 0.18
        self.corners = [
            {"x": margin_x, "y": margin_y},
            {"x": w - margin_x, "y": margin_y},
            {"x": w - margin_x, "y": h - margin_y},
            {"x": margin_x, "y": h - margin_y},
        ]
        self.reset_insets()
        self.warped = None
        self.display_mode = "photo"
        self.status_var.set(f"Loaded {self.image_path.name}. Detecting display...")
        try:
            self.corners = detect_display(self.image)
            self.warped = warp_screen(self.image, self.corners)
            self.display_mode = "flat"
            self.status_var.set(f"Loaded {self.image_path.name}. Display detected; fine tune corners if needed.")
        except Exception as exc:
            self.status_var.set(f"Loaded {self.image_path.name}. Auto detect failed: {exc}. Drag corners manually.")
        self.render_canvas()

    def auto_detect(self) -> None:
        if self.image is None:
            messagebox.showinfo("Choose image", "Choose an image first.")
            return
        try:
            self.corners = detect_display(self.image)
            self.warped = warp_screen(self.image, self.corners)
            self.reset_insets()
            self.display_mode = "flat"
            self.status_var.set("Display detected. Fine tune corners if needed, then Analyze Grid.")
            self.render_canvas()
        except Exception as exc:
            messagebox.showerror("Auto detect failed", str(exc))

    def set_mode(self, mode: str) -> None:
        self.display_mode = mode
        if mode == "flat" and self.image is not None and self.corners:
            self.warped = warp_screen(self.image, self.corners)
        self.render_canvas()

    def get_insets(self) -> dict[str, float]:
        return {key: max(0.0, min(0.25, var.get() / 100.0)) for key, var in self.inset_vars.items()}

    def reset_insets(self) -> None:
        for var in self.inset_vars.values():
            var.set(0.0)

    def get_grid_corners(self) -> list[dict[str, float]]:
        if len(self.corners) != 4:
            return []
        insets = self.get_insets()
        u1 = insets["left"]
        u2 = 1.0 - insets["right"]
        v1 = insets["top"]
        v2 = 1.0 - insets["bottom"]
        return [
            bilinear(self.corners, u1, v1),
            bilinear(self.corners, u2, v1),
            bilinear(self.corners, u2, v2),
            bilinear(self.corners, u1, v2),
        ]

    def canvas_to_image_point(self, x: float, y: float) -> dict[str, float]:
        return {"x": (x - self.offset_x) / self.scale, "y": (y - self.offset_y) / self.scale}

    def image_to_canvas_point(self, point: dict[str, float]) -> tuple[float, float]:
        return point["x"] * self.scale + self.offset_x, point["y"] * self.scale + self.offset_y

    def render_canvas(self) -> None:
        self.canvas.delete("all")
        source = self.image
        if self.display_mode == "flat":
            source = self.warped
        if source is None:
            self.canvas.create_text(20, 20, anchor="nw", fill="white", text="Choose an MCMDU image.")
            return
        canvas_w = max(1, self.canvas.winfo_width())
        canvas_h = max(1, self.canvas.winfo_height())
        self.scale = min(canvas_w / source.width, canvas_h / source.height)
        self.offset_x = int((canvas_w - source.width * self.scale) * 0.5)
        self.offset_y = int((canvas_h - source.height * self.scale) * 0.5)
        shown = source.resize((max(1, int(source.width * self.scale)), max(1, int(source.height * self.scale))))
        self.photo_image = ImageTk.PhotoImage(shown)
        self.canvas.create_image(self.offset_x, self.offset_y, anchor="nw", image=self.photo_image)
        if self.display_mode == "photo":
            self.draw_photo_grid()
            self.draw_corners()
        else:
            self.draw_grid_overlay(source.width, source.height)

    def draw_photo_grid(self) -> None:
        grid_corners = self.get_grid_corners()
        if not grid_corners:
            return
        points = [self.image_to_canvas_point(point) for point in grid_corners]

        def point_at(u: float, v: float) -> dict[str, float]:
            return bilinear(
                [{"x": x, "y": y} for x, y in points],
                u,
                v,
            )

        for col in range(COLS + 1):
            u = col / COLS
            top = point_at(u, 0)
            bottom = point_at(u, 1)
            self.canvas.create_line(top["x"], top["y"], bottom["x"], bottom["y"], fill="#189a8d")
        for row in range(ROWS + 1):
            v = row / ROWS
            left = point_at(0, v)
            right = point_at(1, v)
            self.canvas.create_line(left["x"], left["y"], right["x"], right["y"], fill="#189a8d")

    def draw_corners(self) -> None:
        if not self.corners:
            return
        points = [self.image_to_canvas_point(point) for point in self.corners]
        flat_points = [coord for point in points for coord in point]
        self.canvas.create_polygon(*flat_points, outline="#00d7ff", fill="", width=2)
        labels = ["TL", "TR", "BR", "BL"]
        for index, (x, y) in enumerate(points):
            self.canvas.create_oval(x - 7, y - 7, x + 7, y + 7, fill="#00d7ff", outline="white")
            self.canvas.create_text(x + 12, y, anchor="w", fill="white", text=labels[index])

    def draw_grid_overlay(self, image_w: int, image_h: int) -> None:
        insets = self.get_insets()
        left = self.offset_x + image_w * self.scale * insets["left"]
        right = self.offset_x + image_w * self.scale * (1.0 - insets["right"])
        top = self.offset_y + image_h * self.scale * insets["top"]
        bottom = self.offset_y + image_h * self.scale * (1.0 - insets["bottom"])
        grid_w = right - left
        grid_h = bottom - top
        for col in range(COLS + 1):
            x = left + (grid_w * col) / COLS
            color = "#00d7ff" if col in (0, COLS) else "#2b8a9a"
            self.canvas.create_line(x, top, x, bottom, fill=color)
        for row in range(ROWS + 1):
            y = top + (grid_h * row) / ROWS
            color = "#00d7ff" if row in (0, ROWS) else "#2b8a9a"
            self.canvas.create_line(left, y, right, y, fill=color)
        for row in range(ROWS):
            y = top + (grid_h * (row + 0.5)) / ROWS
            self.canvas.create_text(max(10, left - 16), y, fill="white", text=str(row + 1), font=("Arial", 10))
        for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
            x = left + (grid_w * (col + 0.5)) / COLS
            self.canvas.create_text(x, max(10, top - 14), fill="white", text=str(col), font=("Arial", 9))

    def on_press(self, event: Any) -> None:
        if self.display_mode != "photo" or not self.corners:
            return
        distances = []
        for index, corner in enumerate(self.corners):
            x, y = self.image_to_canvas_point(corner)
            distances.append((math.hypot(event.x - x, event.y - y), index))
        distance, index = min(distances)
        if distance <= 25:
            self.drag_corner = index

    def on_drag(self, event: Any) -> None:
        if self.drag_corner is None or self.image is None:
            return
        point = self.canvas_to_image_point(event.x, event.y)
        point["x"] = max(0, min(self.image.width, point["x"]))
        point["y"] = max(0, min(self.image.height, point["y"]))
        self.corners[self.drag_corner] = point
        self.warped = None
        self.render_canvas()

    def on_release(self, _event: Any) -> None:
        self.drag_corner = None

    def analyze(self) -> None:
        if self.image is None or len(self.corners) != 4:
            messagebox.showinfo("Choose image", "Choose an image and align the four screen corners first.")
            return
        try:
            grid_corners = self.get_grid_corners()
            grid, _warped, box_count, word_count = analyze_image(self.image, grid_corners)
            self.warped = warp_screen(self.image, self.corners)
            self.source_grid = grid
            self.set_grid(grid)
            self.display_mode = "flat"
            self.render_canvas()
            self.status_var.set(f"OCR complete: {box_count} character boxes, {word_count} text blocks.")
        except Exception as exc:
            messagebox.showerror("Analyze failed", str(exc))

    def remember_corrections(self) -> None:
        current = self.get_grid()
        corrections = load_corrections()
        learned = 0
        for row_index in range(ROWS):
            original = re.sub(r"\s+", " ", "".join(cell or " " for cell in self.source_grid[row_index])).strip()
            corrected = "".join(cell or " " for cell in current[row_index])[:COLS].ljust(COLS)
            corrected_list = list(corrected)
            corrected_list[0] = " "
            corrected_list[COLS - 1] = " "
            corrected = "".join(corrected_list)
            if original and corrected.strip() and original != re.sub(r"\s+", " ", corrected).strip():
                corrections[original] = corrected
                learned += 1
        save_corrections(corrections)
        self.source_grid = current
        self.status_var.set(f"Remembered {learned} corrected row pattern(s). Total corrections: {len(corrections)}.")

    def export_word(self) -> None:
        grid = self.get_grid()
        default_name = f"mcmdu-grid-{uuid.uuid4().hex[:8]}.docx"
        filename = filedialog.asksaveasfilename(
            title="Save Word File",
            initialdir=str(EXPORTS),
            initialfile=default_name,
            defaultextension=".docx",
            filetypes=[("Word document", "*.docx")],
        )
        if not filename:
            return
        try:
            export_docx(grid, Path(filename))
            self.status_var.set(f"Exported Word file: {filename}")
        except Exception as exc:
            messagebox.showerror("Export failed", str(exc))


def main() -> None:
    ensure_dirs()
    root = Tk()
    root.geometry("1380x820")
    McmduDesktopTool(root)
    root.mainloop()


if __name__ == "__main__":
    main()
