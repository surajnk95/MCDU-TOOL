from __future__ import annotations

import base64
import csv
import difflib
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import threading
import time
from concurrent.futures import ThreadPoolExecutor
import uuid
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import math

import numpy as np
from PIL import Image, ImageDraw, ImageEnhance, ImageFilter, ImageFont, ImageOps


ROOT = Path(__file__).resolve().parent
STATIC = ROOT / "static"
DATA = ROOT / "data"
EXPORTS = DATA / "exports"
CORRECTIONS = DATA / "corrections.json"
TEMPLATES = DATA / "templates.json"
BLANK_TEMPLATE_KEY = "__BLANK__"
VOCAB_WORDS = DATA / "mcdu_user_words.txt"
VOCAB_PATTERNS = DATA / "mcdu_user_patterns.txt"
# Tesseract user-patterns syntax: \n matches a digit, other characters are
# literal. (NOT regex — \d / \. are rejected as "Invalid user pattern".)
MCDU_TESSERACT_PATTERNS = [
    r"FL\n\n",
    r"FL\n\n\n",
    r".\n\n\n",
    r"\n\n\n\n",
    r"\n\n\n\n\n",
    r"\n\n/\n\n",
    r"\n\n\n/\n\n\n",
]

ROWS = 13
COLS = 40
FIRST_DATA_COL = 1
LAST_DATA_COL = 38
SCREEN_W = 1600
MIN_SCREEN_H = 900
MAX_SCREEN_H = 1400
OCR_WHITELIST = "ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789°º˚/.-<>%:"

# B1: Candidate monospace font paths for the glyph atlas, tried in order.
_ATLAS_FONT_PATHS: list[str] = [
    # macOS
    "/System/Library/Fonts/Supplemental/Andale Mono.ttf",
    "/System/Library/Fonts/Supplemental/Courier New.ttf",
    "/System/Library/Fonts/Supplemental/Courier New Bold.ttf",
    "/System/Library/Fonts/Supplemental/PTMono.ttc",
    "/System/Library/Fonts/Monaco.ttf",
    # Linux
    "/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationMono-Bold.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSansMono-Bold.ttf",
    "/usr/share/fonts/truetype/freefont/FreeMono.ttf",
    "/usr/share/fonts/truetype/ubuntu/UbuntuMono-R.ttf",
    # Windows
    "C:/Windows/Fonts/cour.ttf",
    "C:/Windows/Fonts/courbd.ttf",
    "C:/Windows/Fonts/consola.ttf",
    "C:/Windows/Fonts/lucon.ttf",
]
# B1: atlas classification on/off.  Off because cross-font accuracy caps at
# ~75 % on SF Mono (the only available non-atlas system monospace), which is
# below the 85 % threshold required for production use.
ATLAS_ENABLED: bool = False
# B1: canonical glyph patch dimensions for bbox-normalised atlas templates.
_ATLAS_CANON_H: int = 28   # target glyph-content height inside template
_ATLAS_CANON_W: int = 20   # target glyph-content width inside template
_ATLAS_PAD_Y: int = 6      # padding above/below canonical region
_ATLAS_PAD_X: int = 5      # padding left/right of canonical region
_ATLAS_TMPL_H: int = _ATLAS_CANON_H + 2 * _ATLAS_PAD_Y   # 40
_ATLAS_TMPL_W: int = _ATLAS_CANON_W + 2 * _ATLAS_PAD_X   # 30
_ATLAS_THRESHOLD: float = 0.55   # minimum NCC score to accept a classification
_ATLAS_MARGIN: float = 0.08      # minimum gap between 1st and 2nd-best scores
_ATLAS_MAX_CONFIDENCE: float = 0.55  # cap on confidence written into grid
MAX_REQUEST_BYTES = 50 * 1024 * 1024
MAX_EXPORT_FILES = 80
BLUR_THRESHOLD = 80.0         # variance-of-Laplacian below this → blurry warning (#25)
BLUR_WARNING_THRESHOLD = 150.0  # between BLUR_THRESHOLD and this → marginal quality (#D3)
_TSV_EARLY_WINNER_SCORE = 55.0  # C2: stop extra PSMs when a variant already scores this well
_TSV_GIVEUP_SCORE = 5.0         # C2: give up on a variant after 2 PSMs if score this poor
_STRIP_SKIP_CONFIDENCE = 0.72   # C3: skip per-row strip when mean filled-cell confidence >= this
EXPORT_MAX_AGE_SECONDS = 7 * 24 * 60 * 60
MCDU_WORD_HINTS = {
    "ACT",
    "RTE",
    "LEGS",
    "CRZ",
    "ALT",
    "SPD",
    "LRC",
    "D/D",
    "ECON",
    "RTA",
    "STEP",
    "ENG",
    "OUT",
    "FUEL",
    "RECMD",
    "MAX",
    "OPT",
    "TO",
    "SEL",
    "FL",
}
MCDU_VOCABULARY = {
    "ACT",
    "ALT",
    "ANRU",
    "CRZ",
    "DATA",
    "DES",
    "ECON",
    "ENG",
    "FUEL",
    "GPS",
    "INDEX",
    "IRU",
    "LEGS",
    "LRC",
    "MAX",
    "NAV",
    "NM",
    "OFF",
    "ON",
    "OPT",
    "OUT",
    "POS",
    "RADIO",
    "RECMD",
    "REF",
    "RTA",
    "RTE",
    "SEL",
    "SELECT",
    "SENSOR",
    "SPD",
    "STEP",
    "TO",
    "TRUE",
}
MCDU_PHRASE_REPLACEMENTS = (
    ("ACTRTACRZ", "ACT RTA CRZ"),
    ("ACTLRCD/D", "ACT LRC D/D"),
    ("MODLRCD/D", "MOD LRC D/D"),
    ("RTAPROGRESS", "RTA PROGRESS"),
    ("DESFORECAST", "DES FORECAST"),
    ("OFFPATHDES", "OFFPATH DES"),
    ("ENGOUT", "ENG OUT"),
    ("ALLENG", "ALL ENG"),
    ("CRZALT", "CRZ ALT"),
    ("RTASPD", "RTA SPD"),
    ("SELSPD", "SEL SPD"),
    ("RECHD", "RECMD"),
    ("NAX", "MAX"),
    ("TA/PUEL", "ETA/FUEL"),
    ("TA/FUEL", "ETA/FUEL"),
    ("PUEL", "FUEL"),
    ("KOFIETA/FUEL", "KBFI ETA/FUEL"),
    ("ETA/FUEL", "ETA/FUEL"),
    ("ETAFUEL", "ETA/FUEL"),
    ("LRCSPD", "LRC SPD"),  # B3: enables fuzzy correction of one-char variants (LRCSPO etc.)
)
# Closed vocabulary of MCDU page titles used by snap_title_row (A5).
# Variable fields (M.xxx, xxxKT, RTE #) are handled separately with regex.
MCDU_TITLE_VOCABULARY: tuple[str, ...] = (
    "ACT RTA CRZ",
    "MOD RTA CRZ",
    "ACT LRC D/D",
    "MOD LRC D/D",
    "ACT ECON CRZ",
    "MOD ECON CRZ",
    "ACT ECON D/D",
    "MOD ECON D/D",
    "RTA PROGRESS",
    "DES FORECAST",
    "OFFPATH DES",
)


def find_tesseract() -> str:
    configured = os.environ.get("TESSERACT_CMD")
    if configured:
        return configured
    found = shutil.which("tesseract")
    if found:
        return found
    candidates = [
        Path("/opt/homebrew/bin/tesseract"),
        Path("/usr/local/bin/tesseract"),
        Path("C:/Program Files/Tesseract-OCR/tesseract.exe"),
        Path("C:/Program Files (x86)/Tesseract-OCR/tesseract.exe"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return str(candidate)
    return "tesseract"


TESSERACT = find_tesseract()
_PADDLE_OCR: Any | None = None
_PADDLE_ERROR = ""
_PADDLE_LOCK = threading.Lock()
_PADDLE_READY = threading.Event()
_PADDLE_INITIALIZING = False


def ensure_dirs() -> None:
    DATA.mkdir(exist_ok=True)
    EXPORTS.mkdir(exist_ok=True)
    if not CORRECTIONS.exists():
        CORRECTIONS.write_text("{}", encoding="utf-8")
    if not TEMPLATES.exists():
        TEMPLATES.write_text("{}", encoding="utf-8")
    VOCAB_WORDS.write_text("\n".join(sorted(MCDU_VOCABULARY)), encoding="utf-8")
    VOCAB_PATTERNS.write_text("\n".join(MCDU_TESSERACT_PATTERNS), encoding="utf-8")
    cleanup_exports()


def cleanup_exports() -> None:
    if not EXPORTS.exists():
        return
    now = time.time()
    for suffix in (".png", ".docx"):  # D1: also age out exported Word files
        files = sorted(
            (path for path in EXPORTS.iterdir() if path.is_file() and path.suffix.lower() == suffix),
            key=lambda path: path.stat().st_mtime,
            reverse=True,
        )
        for index, path in enumerate(files):
            try:
                if index >= MAX_EXPORT_FILES or now - path.stat().st_mtime > EXPORT_MAX_AGE_SECONDS:
                    path.unlink()
            except OSError:
                continue


def json_response(handler: SimpleHTTPRequestHandler, status: int, payload: dict[str, Any]) -> None:
    body = json.dumps(payload).encode("utf-8")
    handler.send_response(status)
    handler.send_header("Content-Type", "application/json")
    handler.send_header("Content-Length", str(len(body)))
    handler.end_headers()
    handler.wfile.write(body)


def read_json_body(handler: SimpleHTTPRequestHandler) -> dict[str, Any]:
    length = int(handler.headers.get("Content-Length", "0"))
    if length <= 0:
        return {}
    if length > MAX_REQUEST_BYTES:
        raise ValueError("Image request is too large. Use an image below 50 MB.")
    raw = handler.rfile.read(length)
    return json.loads(raw.decode("utf-8"))


def load_image(data_url: str) -> Image.Image:
    if "," in data_url:
        data_url = data_url.split(",", 1)[1]
    raw = base64.b64decode(data_url)
    return ImageOps.exif_transpose(Image.open(io.BytesIO(raw))).convert("RGB")


def edge_length(a: dict[str, float], b: dict[str, float]) -> float:
    return math.hypot(a["x"] - b["x"], a["y"] - b["y"])


def screen_size_from_corners(corners: list[dict[str, float]]) -> tuple[int, int]:
    top = edge_length(corners[0], corners[1])
    right = edge_length(corners[1], corners[2])
    bottom = edge_length(corners[2], corners[3])
    left = edge_length(corners[3], corners[0])
    width = max(1.0, (top + bottom) * 0.5)
    height = max(1.0, (left + right) * 0.5)
    aspect = width / height
    normalized_height = int(round(SCREEN_W / aspect))
    return SCREEN_W, max(MIN_SCREEN_H, min(MAX_SCREEN_H, normalized_height))


def solve_linear_system(matrix: list[list[float]], values: list[float]) -> list[float]:
    n = len(values)
    augmented = [row[:] + [values[i]] for i, row in enumerate(matrix)]

    for pivot in range(n):
        max_row = max(range(pivot, n), key=lambda row: abs(augmented[row][pivot]))
        augmented[pivot], augmented[max_row] = augmented[max_row], augmented[pivot]
        pivot_value = augmented[pivot][pivot]
        if abs(pivot_value) < 1e-12:
            raise ValueError("Screen corners are too close together to calculate perspective.")

        for col in range(pivot, n + 1):
            augmented[pivot][col] /= pivot_value

        for row in range(n):
            if row == pivot:
                continue
            factor = augmented[row][pivot]
            for col in range(pivot, n + 1):
                augmented[row][col] -= factor * augmented[pivot][col]

    return [augmented[row][n] for row in range(n)]


def perspective_coefficients(src: list[dict[str, float]], dst_size: tuple[int, int]) -> list[float]:
    width, height = dst_size
    dst = [
        {"x": 0.0, "y": 0.0},
        {"x": float(width), "y": 0.0},
        {"x": float(width), "y": float(height)},
        {"x": 0.0, "y": float(height)},
    ]

    matrix: list[list[float]] = []
    values: list[float] = []
    for d, s in zip(dst, src):
        x = d["x"]
        y = d["y"]
        u = s["x"]
        v = s["y"]
        matrix.append([x, y, 1, 0, 0, 0, -u * x, -u * y])
        values.append(u)
        matrix.append([0, 0, 0, x, y, 1, -v * x, -v * y])
        values.append(v)
    return solve_linear_system(matrix, values)


def estimate_grid_origin(image: Image.Image) -> tuple[float, float, float, float]:
    """Find the (origin, scale) pair per axis minimising ink at internal grid boundaries.

    Returns (origin_x, origin_y, scale_x, scale_y).  Scale defaults to 1.0 when
    evidence for a pitch correction is weak; origin defaults to 0.0 when evidence
    for a phase correction is weak.  A correctly-pitched image produces
    scale == 1.0 so that the caller can apply a byte-identical fast path.
    """
    gray = np.asarray(ImageOps.grayscale(image), dtype=np.float32)
    if gray.size == 0:
        return 0.0, 0.0, 1.0, 1.0
    threshold = max(118.0, float(np.percentile(gray, 86)))
    ink = gray >= threshold
    if float(np.mean(ink)) < 0.002:
        return 0.0, 0.0, 1.0, 1.0

    height, width = ink.shape

    def best_axis_origin_and_scale(axis: int, length: int, count: int) -> tuple[float, float]:
        """Return (origin, scale) for one axis."""
        nominal_pitch = length / count
        # Projected ink profile: mean ink fraction at each position along this axis,
        # averaged over the middle 96% of the perpendicular axis to avoid edge artefacts.
        start_other = int((width if axis == 0 else height) * 0.02)
        end_other = int((width if axis == 0 else height) * 0.98)
        if axis == 0:
            profile = ink[start_other:end_other, :].mean(axis=0).astype(np.float64)
        else:
            profile = ink[:, start_other:end_other].mean(axis=1).astype(np.float64)
        # Cumulative sum enables O(1) strip-mean lookups: mean(profile[lo:hi]) = (cum[hi]-cum[lo])/(hi-lo)
        cum = np.concatenate([[0.0], np.cumsum(profile)])  # shape (length+1,)

        def raw_ink_scores(pitch: float, origins: np.ndarray) -> np.ndarray:
            """Mean boundary ink fraction for each origin, vectorised over all origins at once."""
            half_strip = max(1.0, pitch * 0.035)
            totals = np.zeros(len(origins), dtype=np.float64)
            for boundary in range(1, count):
                positions = origins + boundary * pitch
                low_arr = np.clip(np.floor(positions - half_strip).astype(int), 0, length)
                high_arr = np.clip(np.ceil(positions + half_strip).astype(int) + 1, 0, length)
                valid = high_arr > low_arr
                denom = np.where(valid, high_arr - low_arr, 1)
                strip_means = np.where(
                    valid, (cum[high_arr] - cum[low_arr]) / denom, 0.0
                )
                totals += strip_means
            return totals / max(1, count - 1)

        phase_candidates = np.linspace(-nominal_pitch * 0.45, nominal_pitch * 0.45, 145)
        scale_candidates = np.linspace(0.97, 1.03, 13)

        # Baseline: best achievable raw ink score at scale == 1.0 over all phase candidates.
        nominal_raw = raw_ink_scores(nominal_pitch, phase_candidates)
        best_nominal_raw = float(np.min(nominal_raw))

        # Search non-unit scales for a strictly better (origin, scale) pair.
        best_scale = 1.0
        best_origin_nonu = 0.0
        best_raw_nonu = best_nominal_raw  # start equal; only improve if strictly lower

        for sc in scale_candidates:
            if abs(sc - 1.0) < 1e-9:
                continue
            pitch = nominal_pitch * sc
            origins = np.linspace(-pitch * 0.45, pitch * 0.45, 145)
            raw = raw_ink_scores(pitch, origins)
            idx = int(np.argmin(raw))
            if float(raw[idx]) < best_raw_nonu:
                best_raw_nonu = float(raw[idx])
                best_scale = float(sc)
                best_origin_nonu = float(origins[idx])

        # Safety: only deviate from scale==1.0 when evidence is strong (>6% ink reduction).
        if best_scale != 1.0 and best_raw_nonu < best_nominal_raw * 0.94:
            return best_origin_nonu, best_scale

        # Fall back to scale==1.0 with existing phase safety threshold.
        nominal_with_penalty = nominal_raw + np.abs(phase_candidates) / nominal_pitch * 0.0015
        best_phase_idx = int(np.argmin(nominal_with_penalty))
        zero_idx = int(np.argmin(np.abs(phase_candidates)))
        if nominal_with_penalty[best_phase_idx] >= nominal_with_penalty[zero_idx] * 0.94:
            return 0.0, 1.0
        return float(phase_candidates[best_phase_idx]), 1.0

    ox, sx = best_axis_origin_and_scale(0, width, COLS)
    oy, sy = best_axis_origin_and_scale(1, height, ROWS)
    return ox, oy, sx, sy


def align_warp_to_grid(image: Image.Image) -> tuple[Image.Image, tuple[float, float]]:
    origin_x, origin_y, scale_x, scale_y = estimate_grid_origin(image)
    scale_is_unity = abs(scale_x - 1.0) < 1e-9 and abs(scale_y - 1.0) < 1e-9
    if abs(origin_x) < 0.1 and abs(origin_y) < 0.1 and scale_is_unity:
        return image, (0.0, 0.0)
    # PIL AFFINE maps destination→source: src_x = scale_x*dst_x + origin_x.
    # When scale != 1 the pitch is corrected; after the transform the grid sits at
    # its nominal position, so we return (0,0) for the browser grid-alignment offset.
    aligned = image.transform(
        image.size,
        Image.Transform.AFFINE,
        (scale_x, 0.0, origin_x, 0.0, scale_y, origin_y),
        Image.Resampling.BICUBIC,
        fillcolor=(0, 0, 0),
    )
    if not scale_is_unity:
        return aligned, (0.0, 0.0)
    return aligned, (origin_x, origin_y)


def warp_screen_with_alignment(
    image: Image.Image,
    corners: list[dict[str, float]],
) -> tuple[Image.Image, tuple[float, float]]:
    screen_size = screen_size_from_corners(corners)
    coeffs = perspective_coefficients(corners, screen_size)
    warped = image.transform(
        screen_size,
        Image.Transform.PERSPECTIVE,
        coeffs,
        Image.Resampling.BICUBIC,
    )
    return align_warp_to_grid(warped)


def warp_screen(image: Image.Image, corners: list[dict[str, float]]) -> Image.Image:
    warped, _ = warp_screen_with_alignment(image, corners)
    return warped


def max_channel_gray(image: Image.Image) -> Image.Image:
    """Return max(R, G, B) per pixel as an L-mode image (HSV Value channel).

    All MCDU text colors (white, magenta, cyan, amber, green) are bright in at
    least one channel; the dark background is uniformly low.  Using the maximum
    channel preserves every color while keeping the background dark — superior to
    a plain luminance average, which attenuates magenta and amber significantly.
    """
    arr = np.asarray(image.convert("RGB"), dtype=np.uint8)
    return Image.fromarray(np.max(arr, axis=2), mode="L")


def preprocess_for_ocr(image: Image.Image) -> Image.Image:
    gray = max_channel_gray(image)
    gray = ImageOps.invert(gray)
    gray = ImageEnhance.Contrast(gray).enhance(2.8)
    gray = gray.filter(ImageFilter.SHARPEN)
    return gray


def preprocessing_variants(image: Image.Image) -> list[tuple[str, Image.Image]]:
    gray = max_channel_gray(image)
    variants: list[tuple[str, Image.Image]] = [
        ("contrast", preprocess_for_ocr(image)),
        ("inverted", ImageOps.invert(gray)),
        ("grayscale", ImageOps.autocontrast(gray, cutoff=1)),
    ]

    try:
        import cv2  # optional; already in requirements.txt as opencv-python-headless

        gray_arr = np.asarray(gray, dtype=np.uint8)

        # #11 — 2× upscale so Tesseract sees ~30–40 px cap height
        rgb_arr = np.asarray(image.convert("RGB"), dtype=np.uint8)
        h, w = rgb_arr.shape[:2]
        up_arr = cv2.resize(rgb_arr, (w * 2, h * 2), interpolation=cv2.INTER_CUBIC)
        variants.append(("upscale2x", preprocess_for_ocr(Image.fromarray(up_arr, mode="RGB"))))

        # #12 — Denoise (fast NL-means removes phone sensor grain)
        denoised = cv2.fastNlMeansDenoising(gray_arr, None, h=10, templateWindowSize=7, searchWindowSize=21)
        variants.append(("denoised", ImageOps.invert(Image.fromarray(denoised, mode="L"))))

        # #12 — Unsharp mask (counters soft blur from defocus / compression)
        blurred = cv2.GaussianBlur(gray_arr, (0, 0), 2.0)
        sharpened = np.clip(cv2.addWeighted(gray_arr, 1.5, blurred, -0.5, 0), 0, 255).astype(np.uint8)
        variants.append(("unsharp", ImageOps.invert(Image.fromarray(sharpened, mode="L"))))

        # #13 — Otsu binarization (global threshold; works well on bimodal MCDU histograms)
        _, otsu = cv2.threshold(gray_arr, 0, 255, cv2.THRESH_BINARY_INV | cv2.THRESH_OTSU)
        variants.append(("otsu", Image.fromarray(otsu, mode="L")))

        # #13 — Adaptive threshold (Gaussian neighbourhood; handles uneven glare)
        adapted = cv2.adaptiveThreshold(
            gray_arr, 255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV,
            blockSize=21, C=10,
        )
        variants.append(("adaptive", Image.fromarray(adapted, mode="L")))

    except ImportError:
        pass  # OpenCV absent — the three pure-PIL variants above are still returned

    return variants


# ---------------------------------------------------------------------------
# #14 — Magenta cursor erasure
# #15 — Glare suppression
# #16 — Coloured message-box OCR (non-inverted preprocessing)
# #17 — Entry-field outline erasure
# All operate on the perspective-corrected (warped) MCDU image.
# ---------------------------------------------------------------------------


def erase_magenta_cursor(rgb: np.ndarray) -> np.ndarray:
    """Detect and inpaint the magenta + crosshair cursor (#14).

    The cursor is distinguished from magenta text by its cross shape: a
    significant fraction of its pixels lie on both the horizontal AND vertical
    centre strips of the component bounding box.
    """
    try:
        import cv2
    except ImportError:
        return rgb

    hsv = cv2.cvtColor(rgb, cv2.COLOR_RGB2HSV)
    # Magenta: OpenCV hue 130-179 (≈ 260-358° standard), high S and V
    mag_mask = cv2.inRange(
        hsv,
        np.array([130, 80, 100], dtype=np.uint8),
        np.array([179, 255, 255], dtype=np.uint8),
    )
    if not np.any(mag_mask):
        return rgb

    num_labels, labels, stats, _ = cv2.connectedComponentsWithStats(mag_mask, connectivity=8)
    inpaint_mask = np.zeros_like(mag_mask)

    for i in range(1, num_labels):
        area = int(stats[i, cv2.CC_STAT_AREA])
        if area < 20 or area > 8000:
            continue
        bx = int(stats[i, cv2.CC_STAT_LEFT])
        by = int(stats[i, cv2.CC_STAT_TOP])
        bw = int(stats[i, cv2.CC_STAT_WIDTH])
        bh = int(stats[i, cv2.CC_STAT_HEIGHT])
        if bw < 6 or bh < 6:
            continue
        aspect = bw / bh
        if aspect < 0.25 or aspect > 4.0:  # not cross-like aspect ratio
            continue

        # Cross-shape test: both the horizontal AND vertical centre strips must
        # account for at least 20 % of the component pixels each.
        ys, xs = np.nonzero(labels == i)
        cx = bx + bw // 2
        cy = by + bh // 2
        arm_half = max(4, min(bw, bh) // 6)
        in_h = (ys >= cy - arm_half) & (ys <= cy + arm_half)
        in_v = (xs >= cx - arm_half) & (xs <= cx + arm_half)
        h_frac = float(np.sum(in_h)) / area
        v_frac = float(np.sum(in_v)) / area
        # Corner-emptiness test: a real crosshair keeps almost all of its pixels on
        # the two arms, leaving the four corners empty. A centred glyph (digit or
        # letter) fills the corners, so reject it. Without this, magenta values
        # (FL204, .860, 12000, ...) would be silently inpainted away.
        in_arms = float(np.sum(in_h | in_v)) / area
        if h_frac >= 0.20 and v_frac >= 0.20 and in_arms >= 0.85:
            inpaint_mask[labels == i] = 255

    if not np.any(inpaint_mask):
        return rgb

    inpaint_mask = cv2.dilate(inpaint_mask, np.ones((3, 3), np.uint8), iterations=1)
    return cv2.inpaint(rgb, inpaint_mask, inpaintRadius=4, flags=cv2.INPAINT_TELEA)


def suppress_glare(rgb: np.ndarray, cell_w: float, cell_h: float) -> np.ndarray:
    """Inpaint large bright-white glare blooms (#15).

    Glare: very bright (V > 240) and nearly unsaturated (S < 30) blobs larger
    than four character cells.  Small bright areas (OCR text) are left alone.
    """
    try:
        import cv2
    except ImportError:
        return rgb

    hsv = cv2.cvtColor(rgb, cv2.COLOR_RGB2HSV)
    glare_mask = cv2.inRange(
        hsv,
        np.array([0, 0, 240], dtype=np.uint8),
        np.array([179, 30, 255], dtype=np.uint8),
    )
    if not np.any(glare_mask):
        return rgb

    min_area = max(500, int(cell_w * cell_h * 4))
    num_labels, labels, stats, _ = cv2.connectedComponentsWithStats(glare_mask, connectivity=8)
    inpaint_mask = np.zeros_like(glare_mask)
    for i in range(1, num_labels):
        if int(stats[i, cv2.CC_STAT_AREA]) >= min_area:
            inpaint_mask[labels == i] = 255

    if not np.any(inpaint_mask):
        return rgb

    return cv2.inpaint(rgb, inpaint_mask, inpaintRadius=5, flags=cv2.INPAINT_TELEA)


def erase_entry_outlines(rgb: np.ndarray, cell_w: float, cell_h: float) -> np.ndarray:
    """Erase thin rectangular entry-field box outlines (#17).

    Detected as 4-vertex closed contours that are small (1–8 cols × 1–2 rows),
    have a low fill ratio (thin border, mostly empty inside), and do not touch
    the image edges.  Matched pixels are set to black (MCDU background).
    """
    try:
        import cv2
    except ImportError:
        return rgb

    gray = np.max(rgb, axis=2).astype(np.uint8)
    blurred = cv2.GaussianBlur(gray, (3, 3), 0)
    edges = cv2.Canny(blurred, 40, 120)

    contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if not contours:
        return rgb

    h_img, w_img = gray.shape
    erase_mask = np.zeros_like(gray)

    for contour in contours:
        area = cv2.contourArea(contour)
        if area < cell_w * cell_h * 0.05:  # skip tiny noise
            continue
        perimeter = cv2.arcLength(contour, True)
        if perimeter < 1:
            continue
        approx = cv2.approxPolyDP(contour, 0.06 * perimeter, True)
        if len(approx) != 4:
            continue
        bx, by, bw, bh = cv2.boundingRect(contour)
        # Skip contours touching the image border (avoids erasing the display frame)
        if bx <= 2 or by <= 2 or bx + bw >= w_img - 2 or by + bh >= h_img - 2:
            continue
        if bw < cell_w * 0.4 or bh < cell_h * 0.3:
            continue
        if bw > cell_w * 9 or bh > cell_h * 2.5:
            continue
        # Thin-outline test: count actual bright pixels in the bounding box.
        # cv2.contourArea includes the interior hole for a ring contour, so it
        # over-counts; counting max-channel pixels above 30 gives the real fill.
        bright = int(np.sum(gray[by : by + bh, bx : bx + bw] > 30))
        if bright / max(1.0, float(bw * bh)) > 0.35:
            continue
        cv2.drawContours(erase_mask, [contour], -1, 255, thickness=2)

    if not np.any(erase_mask):
        return rgb

    result = rgb.copy()
    result[erase_mask > 0] = 0
    return result


def clean_warped_image(warped: Image.Image) -> Image.Image:
    """Apply cursor removal (#14), glare suppression (#15), and entry-outline
    erasure (#17) to the perspective-corrected MCDU display image before OCR.

    Operations are applied in order: outlines → cursor → glare.
    Returns the original image unchanged if OpenCV is not available.
    """
    try:
        import cv2  # noqa: F401
    except ImportError:
        return warped

    rgb = np.asarray(warped.convert("RGB"), dtype=np.uint8)
    h, w = rgb.shape[:2]
    cell_w = w / COLS
    cell_h = h / ROWS

    rgb = erase_entry_outlines(rgb, cell_w, cell_h)  # #17 — remove box outlines first
    rgb = erase_magenta_cursor(rgb)                   # #14 — inpaint cursor
    rgb = suppress_glare(rgb, cell_w, cell_h)         # #15 — inpaint bright blooms

    return Image.fromarray(rgb, mode="RGB")


def _preprocess_message_box_region(crop_rgb: np.ndarray) -> Image.Image:
    """Preprocess a coloured-background message-box crop for Tesseract (#16).

    Uses (255 − S) × V / 255 so white text (S≈0, V≈255) → bright (then
    inverted to black) and the coloured background (S≈200, V≈150) → dim
    (inverted to light grey).  Result: dark text on light background.
    """
    import cv2
    hsv = cv2.cvtColor(crop_rgb, cv2.COLOR_RGB2HSV).astype(np.float32)
    s, v = hsv[:, :, 1], hsv[:, :, 2]
    text_gray = np.clip((255.0 - s) * v / 255.0, 0, 255).astype(np.uint8)
    return Image.fromarray(255 - text_gray, mode="L")


def ocr_message_boxes(warped: Image.Image, geometry: dict) -> list[dict]:
    """Detect coloured message boxes and OCR them without global inversion (#16).

    Bright, saturated filled rectangles (blue/cyan/amber background with white
    text) are cropped, preprocessed with white-text extraction, and passed to
    Tesseract individually.  Returned words carry coordinates in the warped-image
    space so they can be fed directly to place_word().
    """
    try:
        import cv2
    except ImportError:
        return []

    rgb = np.asarray(warped.convert("RGB"), dtype=np.uint8)
    h, w = rgb.shape[:2]
    cell_w = w / COLS
    cell_h = h / ROWS

    hsv = cv2.cvtColor(rgb, cv2.COLOR_RGB2HSV)
    # Saturated (S > 60) and bright (V > 80) — covers blue, cyan, and amber boxes
    color_mask = cv2.inRange(
        hsv,
        np.array([0, 60, 80], dtype=np.uint8),
        np.array([179, 255, 255], dtype=np.uint8),
    )
    # Skip areas smaller than 3 character cells (avoids per-character colour noise)
    min_area = max(500, int(cell_w * cell_h * 3))
    num_labels, labels, stats, _ = cv2.connectedComponentsWithStats(color_mask, connectivity=8)

    words: list[dict] = []
    for i in range(1, num_labels):
        area = int(stats[i, cv2.CC_STAT_AREA])
        if area < min_area:
            continue
        bx = int(stats[i, cv2.CC_STAT_LEFT])
        by = int(stats[i, cv2.CC_STAT_TOP])
        bw = int(stats[i, cv2.CC_STAT_WIDTH])
        bh = int(stats[i, cv2.CC_STAT_HEIGHT])
        # Solid-fill check: scattered coloured text pixels have low fill
        if area / max(1, bw * bh) < 0.50:
            continue
        crop = rgb[by : by + bh, bx : bx + bw]
        if crop.size == 0:
            continue
        preprocessed = _preprocess_message_box_region(crop)
        padded = ImageOps.expand(preprocessed, border=4, fill=255)
        for word in run_tesseract_tsv(padded):
            word["left"] = int(word.get("left", 0)) - 4 + bx
            word["top"] = int(word.get("top", 0)) - 4 + by
            words.append(word)
    return words


def _normalize_cell_patch(arr: np.ndarray) -> list[float]:
    """Center the tight glyph bbox and apply zero-mean/unit-variance normalization.

    Input: 16×24 float32 array (bright glyph on dark background).
    Output: 384-element feature vector matching the format used by
    cell_feature and classify_from_templates / feature_distance.
    Used by both cell_feature (live OCR) and _build_glyph_atlas (B1).
    """
    arr = arr.astype(np.float32)
    try:
        import cv2  # noqa: F401
        threshold_val = max(90.0, float(arr.mean() + arr.std() * 0.65))
        mask = arr > threshold_val
        if mask.any():
            row_idx, col_idx = np.nonzero(mask)
            r0, r1 = int(row_idx.min()), int(row_idx.max())
            c0, c1 = int(col_idx.min()), int(col_idx.max())
            glyph_h = r1 - r0 + 1
            glyph_w = c1 - c0 + 1
            centered = np.zeros((24, 16), dtype=np.float32)
            dr = max(0, (24 - glyph_h) // 2)
            dc = max(0, (16 - glyph_w) // 2)
            src_h = min(glyph_h, 24 - dr)
            src_w = min(glyph_w, 16 - dc)
            centered[dr : dr + src_h, dc : dc + src_w] = arr[r0 : r0 + src_h, c0 : c0 + src_w]
            arr = centered
    except ImportError:
        pass
    mean = float(arr.mean())
    std = float(arr.std())
    if std > 1e-6:
        arr = (arr - mean) / std
    else:
        arr = arr - mean
    return arr.reshape(-1).tolist()


def cell_feature(image: Image.Image, row: int, col: int) -> list[float]:
    """Extract a centered, normalized glyph patch for the given grid cell (#21).

    The glyph's tight bounding box is detected, translated to the centre of the
    16×24 canvas, then the patch is zero-mean / unit-variance normalized so that
    cosine similarity (feature_distance) is scale-invariant.  Falls back to a
    simpler mean-subtraction when OpenCV is absent or the cell is blank.
    """
    screen_w, screen_h = image.size
    cell_w = screen_w / COLS
    cell_h = screen_h / ROWS
    pad_x = cell_w * 0.08
    pad_y = cell_h * 0.08
    box = (
        int(max(0, col * cell_w - pad_x)),
        int(max(0, row * cell_h - pad_y)),
        int(min(screen_w, (col + 1) * cell_w + pad_x)),
        int(min(screen_h, (row + 1) * cell_h + pad_y)),
    )
    crop = max_channel_gray(image.crop(box)).resize((16, 24), Image.Resampling.BILINEAR)
    return _normalize_cell_patch(np.asarray(crop).astype(np.float32))


def feature_distance(a: list[float], b: list[float]) -> float:
    """Cosine-similarity distance ∈ [0, 1]: 0 = identical, 1 = anti-correlated (#21).

    Uses NCC (normalized cross-correlation) so the metric is scale-invariant and
    works correctly regardless of whether the stored template was captured from a
    bright or a dim cell.
    """
    if len(a) != len(b):
        return 1.0
    arr_a = np.asarray(a, dtype=np.float32)
    arr_b = np.asarray(b, dtype=np.float32)
    dot_aa = float(np.dot(arr_a, arr_a))
    dot_bb = float(np.dot(arr_b, arr_b))
    if dot_aa < 1e-12 or dot_bb < 1e-12:
        return 0.0 if dot_aa < 1e-12 and dot_bb < 1e-12 else 1.0
    ncc = float(np.dot(arr_a, arr_b)) / math.sqrt(dot_aa * dot_bb)
    return (1.0 - max(-1.0, min(1.0, ncc))) / 2.0


def load_templates() -> dict[str, list[list[float]]]:
    ensure_dirs()
    try:
        data = json.loads(TEMPLATES.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        data = {}
    templates: dict[str, list[list[float]]] = {}
    for char, features in data.items():
        if not isinstance(char, str) or not char:
            continue
        if not isinstance(features, list):
            continue
        key = char if char == BLANK_TEMPLATE_KEY else char[:1]
        templates[key] = [feature for feature in features if isinstance(feature, list)]
    return templates


def save_templates(templates: dict[str, list[list[float]]]) -> None:
    ensure_dirs()
    compact = {char: features[-30:] for char, features in templates.items() if features}
    TEMPLATES.write_text(json.dumps(compact), encoding="utf-8")


# ─── Glyph Atlas — B1 ───────────────────────────────────────────────────────

_GLYPH_ATLAS: dict[str, list[np.ndarray]] | None = None


def _atlas_bbox_normalize(arr: np.ndarray, threshold: float = 30.0) -> np.ndarray | None:
    """Extract tight glyph bbox, resize to canonical size, return padded uint8 patch.

    Returns a (_ATLAS_TMPL_H × _ATLAS_TMPL_W) uint8 array, or None when the
    input contains no bright pixels.
    """
    mask = arr > threshold
    if not mask.any():
        return None
    rows, cols = np.nonzero(mask)
    r0, r1 = int(rows.min()), int(rows.max())
    c0, c1 = int(cols.min()), int(cols.max())
    glyph = arr[r0 : r1 + 1, c0 : c1 + 1].astype(np.float32)
    gh, gw = glyph.shape
    if gh < 2 or gw < 2:
        return None
    scale = min(_ATLAS_CANON_H / gh, _ATLAS_CANON_W / gw)
    new_h = max(1, int(round(gh * scale)))
    new_w = max(1, int(round(gw * scale)))
    try:
        import cv2 as _cv2
        resized = _cv2.resize(glyph, (new_w, new_h), interpolation=_cv2.INTER_LINEAR)
    except ImportError:
        resized = np.asarray(
            Image.fromarray(glyph).resize((new_w, new_h), Image.Resampling.BILINEAR),
            dtype=np.float32,
        )
    out = np.zeros((_ATLAS_TMPL_H, _ATLAS_TMPL_W), dtype=np.float32)
    y0 = _ATLAS_PAD_Y + (_ATLAS_CANON_H - new_h) // 2
    x0 = _ATLAS_PAD_X + (_ATLAS_CANON_W - new_w) // 2
    out[y0 : y0 + new_h, x0 : x0 + new_w] = resized
    mx = float(out.max())
    if mx > 0:
        out = out * (255.0 / mx)
    return out.astype(np.uint8)


def _atlas_ncc(a: np.ndarray, b: np.ndarray) -> float:
    """Normalized cross-correlation (cosine similarity) for two same-size arrays."""
    af = a.astype(np.float32) - float(a.mean())
    bf = b.astype(np.float32) - float(b.mean())
    na = float(np.linalg.norm(af))
    nb = float(np.linalg.norm(bf))
    if na < 1e-6 or nb < 1e-6:
        return 0.0
    return float(np.dot(af.flatten(), bf.flatten()) / (na * nb))


def _build_glyph_atlas() -> dict[str, list[np.ndarray]]:
    """Render alphanumeric OCR chars as bbox-normalised 2D image patches.

    For each available atlas font, renders each character at 64 pt on a
    200×200 canvas, extracts the tight glyph bounding box, resizes it to
    the canonical (_ATLAS_CANON_H × _ATLAS_CANON_W) representation with
    padding, and stores three variants (clean, blurred, thickened).

    Returns {char: [uint8 ndarray of shape (_ATLAS_TMPL_H, _ATLAS_TMPL_W)]}.
    """
    atlas: dict[str, list[np.ndarray]] = {}
    atlas_chars = [c for c in OCR_WHITELIST if c.isalnum()]

    for path in _ATLAS_FONT_PATHS:
        if not Path(path).exists():
            continue
        try:
            font = ImageFont.truetype(path, 64)
        except Exception:
            continue

        for char in atlas_chars:
            canvas = Image.new("L", (200, 200), 0)
            draw = ImageDraw.Draw(canvas)
            try:
                bb = draw.textbbox((0, 0), char, font=font)
                x = 100 - (bb[2] - bb[0]) // 2 - bb[0]
                y = 100 - (bb[3] - bb[1]) // 2 - bb[1]
                draw.text((x, y), char, fill=255, font=font)
            except Exception:
                continue

            base_arr = np.asarray(canvas, dtype=np.uint8)
            if base_arr.max() < 10:
                continue

            tmpl = _atlas_bbox_normalize(base_arr.astype(np.float32))
            if tmpl is None:
                continue
            blur = _atlas_bbox_normalize(
                np.asarray(
                    Image.fromarray(base_arr).filter(ImageFilter.GaussianBlur(1.5)),
                    dtype=np.float32,
                )
            )
            thick = _atlas_bbox_normalize(
                np.asarray(
                    Image.fromarray(base_arr).filter(ImageFilter.MaxFilter(3)),
                    dtype=np.float32,
                )
            )
            atlas.setdefault(char, []).extend(v for v in [tmpl, blur, thick] if v is not None)

    return atlas


def _get_glyph_atlas() -> dict[str, list[np.ndarray]]:
    """Return the cached glyph atlas, building it once on first call (B1)."""
    global _GLYPH_ATLAS
    if _GLYPH_ATLAS is None:
        _GLYPH_ATLAS = _build_glyph_atlas()
    return _GLYPH_ATLAS


def _atlas_classify_patch(
    cell_arr: np.ndarray,
    atlas: dict[str, list[np.ndarray]],
) -> tuple[str, float] | None:
    """Core atlas classifier: bbox-normalise *cell_arr* then NCC vs each template.

    Returns (char, score) or None when the ink gate fires, the score is below
    _ATLAS_THRESHOLD, or the best-vs-second margin is less than _ATLAS_MARGIN.
    """
    if float(np.mean(cell_arr > 40)) < 0.005:
        return None
    norm = _atlas_bbox_normalize(cell_arr.astype(np.float32))
    if norm is None:
        return None
    scores: dict[str, float] = {}
    for char, templates in atlas.items():
        best = -1.0
        for tmpl in templates:
            s = _atlas_ncc(norm, tmpl)
            if s > best:
                best = s
        scores[char] = best
    if not scores:
        return None
    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    best_char, best_score = ranked[0]
    if best_score < _ATLAS_THRESHOLD:
        return None
    margin = best_score - ranked[1][1] if len(ranked) >= 2 else best_score
    if margin < _ATLAS_MARGIN:
        return None
    return best_char, best_score


def classify_cell_atlas(
    image: Image.Image,
    row: int,
    col: int,
    atlas: dict[str, list[np.ndarray]],
) -> tuple[str, float] | None:
    """Classify a warped-image grid cell against the pre-seeded glyph atlas (B1).

    Extracts the cell crop, applies the ink gate, bbox-normalises the glyph,
    and runs NCC against every atlas template.  Returns (char, score) or None.
    """
    screen_w, screen_h = image.size
    cell_w = screen_w / COLS
    cell_h = screen_h / ROWS
    x0 = int(max(0, col * cell_w))
    y0 = int(max(0, row * cell_h))
    x1 = int(min(screen_w, (col + 1) * cell_w))
    y1 = int(min(screen_h, (row + 1) * cell_h))
    crop = max_channel_gray(image.crop((x0, y0, x1, y1)))
    cell_arr = np.asarray(crop, dtype=np.uint8)
    return _atlas_classify_patch(cell_arr, atlas)


# ────────────────────────────────────────────────────────────────────────────

def classify_from_templates(feature: list[float], templates: dict[str, list[list[float]]]) -> tuple[str, float] | None:
    best_char = ""
    best_distance = 1.0
    for char, features in templates.items():
        for candidate in features:
            distance = feature_distance(feature, candidate)
            if distance < best_distance:
                best_char = char
                best_distance = distance
    if best_char and best_distance <= 0.25:
        return best_char, best_distance
    return None


def word_quality(word: dict[str, Any]) -> float:
    text = clean_ocr_text(str(word.get("text", ""))).upper()
    compact = text.replace(" ", "")
    score = max(0.0, float(word.get("conf") or 0.0))
    score += min(len(compact), 12) * 3.0
    for hint in MCDU_WORD_HINTS:
        if hint in compact:
            score += 35.0
    if re.fullmatch(r"\d+/\d+", compact):
        score += 25.0
    if re.fullmatch(r"FL\d{2,3}", compact):
        score += 30.0
    return score


def ocr_words_score(words: list[dict[str, Any]], image_height: int) -> float:
    score = sum(word_quality(word) for word in words)
    row_count = len(
        {
            int((float(word["top"]) + float(word["height"]) * 0.5) / max(1.0, image_height / ROWS))
            for word in words
        }
    )
    score += row_count * 18.0
    score -= sum(20.0 for word in words if len(str(word["text"])) > 18)
    return score


def initialize_paddle_ocr() -> None:
    global _PADDLE_OCR, _PADDLE_ERROR
    try:
        try:
            from paddleocr import PaddleOCR
        except (ImportError, OSError) as exc:
            _PADDLE_ERROR = f"PaddleOCR unavailable: {exc}"
            return

        constructor_options = (
            {
                "lang": "en",
                "text_detection_model_name": os.environ.get(
                    "PADDLE_DETECTION_MODEL",
                    "PP-OCRv5_mobile_det",
                ),
                "text_recognition_model_name": os.environ.get(
                    "PADDLE_RECOGNITION_MODEL",
                    "en_PP-OCRv5_mobile_rec",
                ),
                "use_doc_orientation_classify": False,
                "use_doc_unwarping": False,
                "use_textline_orientation": False,
            },
            {"lang": "en", "use_angle_cls": False, "show_log": False},
            {"lang": "en"},
        )
        errors: list[str] = []
        for options in constructor_options:
            try:
                _PADDLE_OCR = PaddleOCR(**options)
                return
            except (TypeError, ValueError) as exc:
                errors.append(str(exc))
            except Exception as exc:  # noqa: BLE001 - optional OCR engine boundary
                _PADDLE_ERROR = f"PaddleOCR could not start: {exc}"
                return
        _PADDLE_ERROR = f"PaddleOCR could not start: {'; '.join(errors)}"
    finally:
        _PADDLE_READY.set()


def get_paddle_ocr() -> tuple[Any | None, str]:
    global _PADDLE_INITIALIZING
    if _PADDLE_OCR is not None:
        return _PADDLE_OCR, ""
    if _PADDLE_ERROR:
        return None, _PADDLE_ERROR

    with _PADDLE_LOCK:
        if not _PADDLE_INITIALIZING:
            _PADDLE_INITIALIZING = True
            threading.Thread(
                target=initialize_paddle_ocr,
                name="paddle-ocr-initializer",
                daemon=True,
            ).start()

    wait_seconds = max(0.0, float(os.environ.get("PADDLE_INIT_WAIT_SECONDS", "12")))
    _PADDLE_READY.wait(timeout=wait_seconds)
    if _PADDLE_OCR is not None:
        return _PADDLE_OCR, ""
    if _PADDLE_ERROR:
        return None, _PADDLE_ERROR
    return None, "PaddleOCR models are initializing in the background; analyze again when ready."


def paddle_result_payload(result: Any) -> dict[str, Any] | None:
    if isinstance(result, dict):
        if isinstance(result.get("res"), dict):
            return result["res"]
        if "rec_texts" in result or "dt_polys" in result:
            return result
    json_value = getattr(result, "json", None)
    if callable(json_value):
        json_value = json_value()
    if isinstance(json_value, str):
        try:
            json_value = json.loads(json_value)
        except json.JSONDecodeError:
            json_value = None
    if isinstance(json_value, dict):
        return paddle_result_payload(json_value)
    return None


def paddle_words_from_payload(payload: dict[str, Any]) -> list[dict[str, Any]]:
    def first_present(*keys: str) -> Any:
        for key in keys:
            value = payload.get(key)
            if value is not None:
                return value
        return []

    texts = first_present("rec_texts", "texts")
    scores = first_present("rec_scores", "scores")
    polygons = first_present("dt_polys", "rec_polys", "polys")
    words: list[dict[str, Any]] = []
    for index, text_value in enumerate(texts):
        text = clean_ocr_text(str(text_value))
        if not text or index >= len(polygons):
            continue
        try:
            points = np.asarray(polygons[index], dtype=np.float32).reshape(-1, 2)
        except (TypeError, ValueError):
            continue
        if points.size < 8:
            continue
        left = float(np.min(points[:, 0]))
        top = float(np.min(points[:, 1]))
        right = float(np.max(points[:, 0]))
        bottom = float(np.max(points[:, 1]))
        score = float(scores[index]) if index < len(scores) else 0.5
        if score <= 1.0:
            score *= 100.0
        words.append(
            {
                "text": text,
                "conf": max(0.0, min(100.0, score)),
                "left": int(round(left)),
                "top": int(round(top)),
                "width": max(1, int(round(right - left))),
                "height": max(1, int(round(bottom - top))),
                "engine": "paddle",
            }
        )
    return words


def paddle_words_from_legacy(result: Any) -> list[dict[str, Any]]:
    lines = result
    if isinstance(lines, list) and len(lines) == 1 and isinstance(lines[0], list):
        lines = lines[0]
    words: list[dict[str, Any]] = []
    if not isinstance(lines, list):
        return words
    for line in lines:
        if not isinstance(line, (list, tuple)) or len(line) < 2:
            continue
        polygon, recognition = line[0], line[1]
        if not isinstance(recognition, (list, tuple)) or not recognition:
            continue
        words.extend(
            paddle_words_from_payload(
                {
                    "rec_texts": [recognition[0]],
                    "rec_scores": [recognition[1] if len(recognition) > 1 else 0.5],
                    "dt_polys": [polygon],
                }
            )
        )
    return words


def boxes_overlap_ratio(first: dict[str, Any], second: dict[str, Any]) -> float:
    first_right = float(first["left"]) + float(first["width"])
    first_bottom = float(first["top"]) + float(first["height"])
    second_right = float(second["left"]) + float(second["width"])
    second_bottom = float(second["top"]) + float(second["height"])
    intersection_w = max(0.0, min(first_right, second_right) - max(float(first["left"]), float(second["left"])))
    intersection_h = max(0.0, min(first_bottom, second_bottom) - max(float(first["top"]), float(second["top"])))
    intersection = intersection_w * intersection_h
    first_area = max(1.0, float(first["width"]) * float(first["height"]))
    second_area = max(1.0, float(second["width"]) * float(second["height"]))
    return intersection / min(first_area, second_area)


def merge_edge_words(
    words: list[dict[str, Any]],
    edge_words: list[dict[str, Any]],
    image_size: tuple[int, int],
    padding: int,
) -> list[dict[str, Any]]:
    _, image_h = image_size
    candidates: list[dict[str, Any]] = []
    for raw_word in edge_words:
        word = dict(raw_word)
        word["left"] = int(word["left"]) - padding
        word["top"] = int(word["top"]) - padding
        word["sourcePass"] = "edge"
        bottom = float(word["top"]) + float(word["height"])
        if float(word["top"]) > image_h * 0.10 and bottom < image_h * 0.90:
            continue
        candidates.append(word)
    return merge_unique_words(words, candidates)


def merge_unique_words(
    words: list[dict[str, Any]],
    candidates: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    merged = [dict(word) for word in words]
    for word in candidates:
        if any(boxes_overlap_ratio(word, existing) >= 0.35 for existing in merged):
            continue
        merged.append(dict(word))
    return sorted(merged, key=lambda item: (int(item["top"]), int(item["left"])))


def remap_crop_words(
    words: list[dict[str, Any]],
    crop_origin: tuple[int, int],
    scale: float,
    padding: int,
) -> list[dict[str, Any]]:
    mapped: list[dict[str, Any]] = []
    for raw_word in words:
        word = dict(raw_word)
        word["left"] = int(round((float(word["left"]) - padding) / scale + crop_origin[0]))
        word["top"] = int(round((float(word["top"]) - padding) / scale + crop_origin[1]))
        word["width"] = max(1, int(round(float(word["width"]) / scale)))
        word["height"] = max(1, int(round(float(word["height"]) / scale)))
        word["sourcePass"] = "corner"
        mapped.append(word)
    return mapped


def run_paddle_ocr(image: Image.Image) -> tuple[list[dict[str, Any]], str]:
    engine, error = get_paddle_ocr()
    if engine is None:
        return [], error
    image_array = np.asarray(image.convert("RGB"))
    try:
        if hasattr(engine, "predict"):
            padding = 30
            padded = ImageOps.expand(image, border=padding, fill=(0, 0, 0))
            cell_h = image.height / ROWS
            crop_scale = 2.0
            crop_padding = 40
            focused_regions = [
                (int(image.width * 0.62), 0, image.width, min(image.height, int(cell_h * 1.8))),
                (
                    int(image.width * 0.62),
                    max(0, int(image.height - cell_h * 1.6)),
                    image.width,
                    image.height,
                ),
                (
                    0,
                    int(image.height * 0.52),
                    int(image.width * 0.38),
                    int(image.height * 0.75),
                ),
            ]
            corner_crops: list[tuple[Image.Image, tuple[int, int]]] = []
            for x1, y1, x2, y2 in focused_regions:
                crop = image.crop((x1, y1, x2, y2))
                crop = crop.resize(
                    (max(1, int(crop.width * crop_scale)), max(1, int(crop.height * crop_scale))),
                    Image.Resampling.LANCZOS,
                )
                crop = ImageOps.expand(crop, border=crop_padding, fill=(0, 0, 0))
                corner_crops.append((crop, (x1, y1)))

            inputs = [image_array, np.asarray(padded.convert("RGB"))]
            inputs.extend(np.asarray(crop.convert("RGB")) for crop, _ in corner_crops)
            results = list(engine.predict(input=inputs))
            parsed: list[list[dict[str, Any]]] = []
            for result in results[: len(inputs)]:
                payload = paddle_result_payload(result)
                parsed.append(paddle_words_from_payload(payload) if payload else [])
            words = parsed[0] if parsed else []
            if len(parsed) > 1:
                words = merge_edge_words(words, parsed[1], image.size, padding)
            for index, (_, origin) in enumerate(corner_crops, start=2):
                if index >= len(parsed):
                    continue
                mapped = remap_crop_words(parsed[index], origin, crop_scale, crop_padding)
                words = merge_unique_words(words, mapped)
        else:
            words = paddle_words_from_legacy(engine.ocr(image_array, cls=False))
    except Exception as exc:  # noqa: BLE001 - optional OCR engine boundary
        return [], f"PaddleOCR failed: {exc}"
    return sorted(words, key=lambda item: (int(item["top"]), int(item["left"]))), ""


def _tesseract_extra_args() -> list[str]:
    """Shared Tesseract flags added to every OCR call (#19/#20).

    --oem 1  forces the LSTM engine.
    user_defined_dpi=300  silences DPI warnings and anchors font-size heuristics.
    --user-words / --user-patterns  hint the recognizer toward MCDU vocabulary;
    only included when the files actually exist (written by ensure_dirs).
    """
    args = ["--oem", "1", "-c", "user_defined_dpi=300"]
    if VOCAB_WORDS.exists():
        args.extend(["--user-words", str(VOCAB_WORDS)])
    if VOCAB_PATTERNS.exists():
        args.extend(["--user-patterns", str(VOCAB_PATTERNS)])
    return args


def run_tesseract_tsv(image: Image.Image) -> list[dict[str, Any]]:
    ensure_dirs()
    with tempfile.TemporaryDirectory() as temp_dir:
        source = Path(temp_dir) / "screen.png"
        image.save(source)

        candidates: list[list[dict[str, Any]]] = []
        last_error = ""
        for i, psm in enumerate(("11", "12", "6", "3", "4")):
            command = [
                TESSERACT,
                str(source),
                "stdout",
                "--psm",
                psm,
                "-l",
                "eng",
                "-c",
                f"tessedit_char_whitelist={OCR_WHITELIST}",
                *_tesseract_extra_args(),
                "tsv",
            ]
            try:
                completed = subprocess.run(command, capture_output=True, text=True, check=False)
            except FileNotFoundError as exc:
                raise RuntimeError(
                    "Tesseract OCR was not found. Install Tesseract OCR, or set TESSERACT_CMD to the full tesseract.exe path."
                ) from exc
            if completed.returncode != 0:
                last_error = completed.stderr.strip()
            else:
                rows = csv.DictReader(io.StringIO(completed.stdout), delimiter="\t")
                words: list[dict[str, Any]] = []
                for row in rows:
                    text = clean_ocr_text((row.get("text") or "").strip())
                    if not text:
                        continue
                    try:
                        conf = float(row.get("conf") or -1)
                    except ValueError:
                        conf = -1
                    if conf < 0:
                        continue
                    words.append(
                        {
                            "text": text,
                            "conf": conf,
                            "left": int(row.get("left") or 0),
                            "top": int(row.get("top") or 0),
                            "width": int(row.get("width") or 0),
                            "height": int(row.get("height") or 0),
                            "psm": psm,
                            "engine": "tesseract",
                        }
                    )
                if words:
                    candidates.append(words)
            # C2: after the first two PSM attempts, prune early: stop if the variant
            # already has an excellent result, or has proven to yield nothing useful.
            if i >= 1:
                if candidates:
                    _best_now = max(ocr_words_score(c, image.height) for c in candidates)
                    if _best_now >= _TSV_EARLY_WINNER_SCORE or _best_now < _TSV_GIVEUP_SCORE:
                        break
                else:
                    break  # no output at all after 2+ PSMs — skip remaining
        if not candidates and last_error:
            raise RuntimeError(last_error or "Tesseract OCR failed.")
        if not candidates:
            return []

        best = max(candidates, key=lambda candidate: ocr_words_score(candidate, image.height))
        return sorted(best, key=lambda item: (int(item["top"]), int(item["left"])))


def run_tesseract_boxes(image: Image.Image) -> list[dict[str, Any]]:
    ensure_dirs()
    width, height = image.size
    with tempfile.TemporaryDirectory() as temp_dir:
        source = Path(temp_dir) / "screen.png"
        image.save(source)
        candidates: list[list[dict[str, Any]]] = []
        for psm in ("6", "11"):
            command = [
                TESSERACT,
                str(source),
                "stdout",
                "--psm",
                psm,
                "-l",
                "eng",
                "-c",
                f"tessedit_char_whitelist={OCR_WHITELIST}",
                *_tesseract_extra_args(),
                "makebox",
            ]
            try:
                completed = subprocess.run(command, capture_output=True, text=True, check=False)
            except FileNotFoundError:
                return []
            if completed.returncode != 0:
                continue

            boxes: list[dict[str, Any]] = []
            for line in completed.stdout.splitlines():
                parts = line.split()
                if len(parts) < 5:
                    continue
                char = clean_ocr_text(parts[0])
                if not char:
                    continue
                try:
                    left = int(parts[1])
                    bottom = int(parts[2])
                    right = int(parts[3])
                    top = int(parts[4])
                except ValueError:
                    continue
                boxes.append(
                    {
                        "text": char[:1],
                        "left": left,
                        "top": max(0, height - top),
                        "width": max(0, right - left),
                        "height": max(0, top - bottom),
                        "psm": psm,
                    }
                )
            if boxes:
                candidates.append(boxes)
        if not candidates:
            return []

        def box_score(boxes: list[dict[str, Any]]) -> float:
            valid = [
                box
                for box in boxes
                if 2 <= float(box["width"]) <= width / COLS * 2.1
                and 2 <= float(box["height"]) <= height / ROWS * 1.4
            ]
            rows = {
                int((float(box["top"]) + float(box["height"]) * 0.5) / max(1.0, height / ROWS))
                for box in valid
            }
            return len(valid) + len(rows) * 4.0

        return max(candidates, key=box_score)


def _run_tesseract_single_line(image: Image.Image) -> list[dict[str, Any]]:
    """Run Tesseract --psm 7 (single text line) on a strip image.

    Lighter-weight than run_tesseract_tsv: no multi-PSM scoring, one call.
    Returns [] silently on any error so callers can treat it as a no-op.
    """
    with tempfile.TemporaryDirectory() as temp_dir:
        source = Path(temp_dir) / "strip.png"
        image.save(source)
        command = [
            TESSERACT,
            str(source),
            "stdout",
            "--psm",
            "7",
            "-l",
            "eng",
            "-c",
            f"tessedit_char_whitelist={OCR_WHITELIST}",
            *_tesseract_extra_args(),
            "tsv",
        ]
        try:
            completed = subprocess.run(command, capture_output=True, text=True, check=False, timeout=20)
        except (FileNotFoundError, subprocess.TimeoutExpired):
            return []
        if completed.returncode != 0:
            return []
        words: list[dict[str, Any]] = []
        for row in csv.DictReader(io.StringIO(completed.stdout), delimiter="\t"):
            text = clean_ocr_text((row.get("text") or "").strip())
            if not text:
                continue
            try:
                conf = float(row.get("conf") or -1)
            except ValueError:
                conf = -1
            if conf < 0:
                continue
            words.append(
                {
                    "text": text,
                    "conf": conf,
                    "left": int(row.get("left") or 0),
                    "top": int(row.get("top") or 0),
                    "width": int(row.get("width") or 0),
                    "height": int(row.get("height") or 0),
                }
            )
        return words


def per_row_strip_ocr(
    warped: Image.Image,
    geometry: dict,
) -> tuple[list[list[str]], list[list[float]]]:
    """Recognise each MCDU row as an isolated strip with Tesseract --psm 7 (#18).

    Slices the warped image into 13 individual data-column strips (cols 1–38),
    upscales each 3×, preprocesses, and runs Tesseract in single-line mode.
    Characters are mapped back to their column positions using the known cell
    width — a layout-independent reading that avoids cross-row confusion and
    scales better than the whole-image multi-PSM pass on blurry photos.

    The result is added to word_candidates and competes row-by-row with other
    preprocessing variants via row_candidate_score.
    """
    cell_w = geometry["cell_w"]
    cell_h = geometry["cell_h"]
    origin_x = geometry["origin_x"]
    origin_y = geometry["origin_y"]
    scale = 3
    data_cols = LAST_DATA_COL - FIRST_DATA_COL + 1

    grid = empty_grid()
    confidence_out = empty_confidence_grid()

    # Pre-compute all row strips that pass the ink gate, then OCR them in parallel.
    _strip_jobs: list[tuple[int, Image.Image, float]] = []
    for row in range(ROWS):
        y1 = max(0, int(origin_y + row * cell_h))
        y2 = min(warped.height, int(origin_y + (row + 1) * cell_h))
        x1 = max(0, int(origin_x + FIRST_DATA_COL * cell_w))
        x2 = min(warped.width, int(origin_x + (LAST_DATA_COL + 1) * cell_w))
        if x2 <= x1 or y2 <= y1:
            continue
        strip = warped.crop((x1, y1, x2, y2))
        # Ink gate: skip strips with almost no bright pixels. PSM 7 assumes the
        # input IS a line of text and will hallucinate low-confidence characters
        # on a blank strip, which row_candidate_score could then pick over a
        # genuinely empty row. A blank strip has near-zero ink, so don't OCR it.
        ink = np.asarray(max_channel_gray(strip), dtype=np.float32)
        if ink.size == 0 or float(np.mean(ink > 80)) < 0.004:
            continue
        enlarged = strip.resize(
            (max(1, strip.width * scale), max(1, strip.height * scale)),
            Image.Resampling.LANCZOS,
        )
        processed_strip = preprocess_for_ocr(enlarged)
        col_w_scaled = enlarged.width / max(1, data_cols)
        _strip_jobs.append((row, processed_strip, col_w_scaled))

    def _ocr_one_strip(
        job: tuple[int, Image.Image, float],
    ) -> tuple[int, list[dict[str, Any]], float]:
        _row, _img, _col_w = job
        try:
            return _row, _run_tesseract_single_line(_img), _col_w
        except Exception:  # noqa: BLE001 — supplemental pass; never block the main result
            return _row, [], _col_w

    # C1: run all per-row strip OCR calls in parallel
    with ThreadPoolExecutor(max_workers=min(max(len(_strip_jobs), 1), ROWS)) as _strip_pool:
        _strip_results = list(_strip_pool.map(_ocr_one_strip, _strip_jobs))

    for row, words, col_w_scaled in _strip_results:
        for word in words:
            text = normalize_mcdu_phrase(str(word.get("text", "")))
            compact = text.replace(" ", "")
            if not compact:
                continue
            word_left = float(word.get("left", 0))
            word_width = max(1.0, float(word.get("width", col_w_scaled)))
            conf = max(0.20, min(0.98, float(word.get("conf", 0)) / 100.0))
            box_cols = max(1, int(round(word_width / col_w_scaled)))
            start_col_idx = max(0, int(round(word_left / col_w_scaled)))
            sequence = text if " " in text and len(text) <= box_cols + 2 else compact
            use_projected_spacing = " " not in sequence and box_cols > len(compact) + 1
            max_len = data_cols - start_col_idx

            for index, char in enumerate(sequence[:max_len]):
                if char.isspace():
                    continue
                if use_projected_spacing:
                    center_x = word_left + word_width * ((index + 0.5) / max(1, len(compact)))
                    col_idx = min(data_cols - 1, int(center_x / col_w_scaled))
                else:
                    col_idx = start_col_idx + index
                col = clamp_data_col(FIRST_DATA_COL + max(0, min(data_cols - 1, col_idx)))
                if not grid[row][col]:
                    grid[row][col] = char
                    confidence_out[row][col] = conf

    return grid, confidence_out


def _ocr_confidence_score(image: Image.Image) -> float:
    """Single-PSM Tesseract pass; return sum of (confidence × character-count).

    Used only for orientation probing — one call per rotation candidate, so the
    function intentionally avoids the multi-PSM loop in run_tesseract_tsv.
    """
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "probe.png"
            image.save(source)
            # No char whitelist here: on some Tesseract builds the whitelist
            # forces reported confidence to 0, which would make this probe a
            # no-op. The probe only needs a relative confidence signal, not
            # clean text, so the whitelist is intentionally omitted.
            command = [
                TESSERACT, str(source), "stdout",
                "--psm", "11", "-l", "eng",
                *_tesseract_extra_args(),
                "tsv",
            ]
            completed = subprocess.run(
                command, capture_output=True, text=True, check=False, timeout=12
            )
            if completed.returncode != 0:
                return 0.0
            total = 0.0
            for row in csv.DictReader(io.StringIO(completed.stdout), delimiter="\t"):
                text = (row.get("text") or "").strip()
                try:
                    conf = float(row.get("conf") or -1)
                except ValueError:
                    conf = -1
                if text and conf >= 15:
                    total += conf * len(text)
            return total
    except Exception:  # noqa: BLE001 - probe must never raise
        return 0.0


def probe_orientation(image: Image.Image) -> int:
    """Return 0, 90 or 270 — the upright correction for a sideways phone photo.

    Scores 0/90/180/270° by OCR confidence on a 300 px downscale, but only ever
    auto-applies a 90 or 270 rotation, and only when the winning sideways angle
    clearly beats upright.  A 180° flip is excluded entirely: phones are not held
    upside down, and 180 is a common false positive that would destroy an upright
    image (use the manual rotate button for the rare real case).  Returns 0 when
    in doubt.  Callers apply PIL Image.rotate(result, expand=True).
    """
    # Downscale to ~1000 px with LANCZOS. A more aggressive downscale (or a
    # BILINEAR filter) aliases the thin MCDU text into orientation-dependent
    # noise and makes the probe pick a wrong rotation on a sharp upright photo.
    scale = min(1.0, 1000.0 / max(image.size))
    small = image.resize(
        (max(1, round(image.width * scale)), max(1, round(image.height * scale))),
        Image.Resampling.LANCZOS,
    )
    scores: dict[int, float] = {}
    for angle in (0, 90, 180, 270):
        candidate = small.rotate(angle, expand=True) if angle else small
        scores[angle] = _ocr_confidence_score(preprocess_for_ocr(candidate))

    upright = scores[0]
    best_angle = 90 if scores[90] >= scores[270] else 270
    best = scores[best_angle]
    if best <= 0:
        return 0
    # A genuinely sideways photo reads almost nothing upright, so the winning
    # rotation dominates. Require a clear margin before disturbing the image.
    if upright > 0 and best < upright * 1.6:
        return 0
    return best_angle


def clean_ocr_text(text: str) -> str:
    text = text.replace("|", "I")
    text = text.replace("º", "°").replace("˚", "°").replace("Â°", "°")
    # Heading / position format, e.g. "000°/0.0NM" — recover a degree that OCR
    # either dropped or misread as o/O before the "/<dist>NM" tail.
    text = re.sub(r"(?<=\d{3})[oO](?=/\d+(?:\.\d+)(?:NM)?\b)", "°", text)
    text = re.sub(r"\b(\d{3})(?=/\d+(?:\.\d+)(?:NM)?\b)", r"\1°", text)
    # Course / track headings, e.g. "217°TRK" — the degree sits between a
    # 3-digit value and TRK and is often dropped or read as o/O/0.
    text = re.sub(r"\b(\d{3})[oO0]?(?=TRK\b)", r"\1°", text)
    # Glidepath angle, e.g. "3.00°" — a degree misread as the letter o/O right
    # after a single-digit decimal (digit 0 is left alone as it is ambiguous).
    text = re.sub(r"\b(\d\.\d{2})[oO]\b", r"\1°", text)
    # Standalone course value, e.g. "268°" — a trailing degree read as o/O at a
    # word boundary. Restricted to the letter o/O so real numbers are untouched.
    text = re.sub(r"\b(\d{3})[oO]\b", r"\1°", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_mcdu_phrase(text: str) -> str:
    text = clean_ocr_text(text).upper()
    compact = text.replace(" ", "")
    if re.fullmatch(r"TOFL\d{2,3}", compact):
        return f"TO FL{compact[4:]}"
    if compact == "KBFIETA/FUEL":
        return "KBFI ETA/FUEL"
    if compact == "TOT/D":
        return "TO T/D"
    if compact in {"ECONSPD", "ECONSP0", "ECONSPOD"}:
        return "ECON SPD"
    suffix = ">" if compact.endswith(">") else ""
    core = compact[:-1] if suffix else compact
    for source, replacement in MCDU_PHRASE_REPLACEMENTS:
        if core == source:
            return f"{replacement}{suffix}"
        if source in core and len(core) <= len(source) + 4:
            return f"{core.replace(source, replacement)}{suffix}"
    # B3: fuzzy phrase normalisation — catches single-character OCR substitutions.
    # Only fires when there is exactly one close match (ratio ≥ 0.80) with a clear
    # margin (≥ 0.10) over the second-best candidate.  Correct text is unchanged.
    _b3_best_r = 0.0
    _b3_best_repl: str | None = None
    _b3_second_r = 0.0
    for _src, _repl in MCDU_PHRASE_REPLACEMENTS:
        _r = difflib.SequenceMatcher(None, core, _src).ratio()
        if _r > _b3_best_r:
            _b3_second_r = _b3_best_r
            _b3_best_r = _r
            _b3_best_repl = _repl
        elif _r > _b3_second_r:
            _b3_second_r = _r
    if _b3_best_repl is not None and _b3_best_r >= 0.80 and _b3_best_r - _b3_second_r >= 0.10:
        return f"{_b3_best_repl}{suffix}"
    for source, replacement in (
        ("RECHD", "RECMD"),
        ("NAX", "MAX"),
        ("PUEL", "FUEL"),
        ("KOFI", "KBFI"),
        ("1RC", "LRC"),
    ):
        if source in core:
            core = core.replace(source, replacement)
    original_core = compact[:-1] if suffix else compact
    if core != original_core:
        return f"{core}{suffix}"
    return text


def mcdu_row_score(text: str) -> float:
    normalized = normalize_mcdu_phrase(text)
    compact = normalized.replace(" ", "")
    score = 0.0
    for word in MCDU_VOCABULARY:
        if word in normalized.split() or word in compact:
            score += 8.0
    for hint in MCDU_WORD_HINTS:
        if hint in compact:
            score += 12.0
    if re.search(r"\d+/\d+", compact):
        score += 10.0
    if re.search(r"FL\d{2,3}", compact):
        score += 10.0
    score += compact.count("<") * 4.0 + compact.count(">") * 4.0
    score += min(len(compact), 20) * 0.4
    score -= len(re.findall(r"(?<![A-Z0-9])[A-Z0-9](?![A-Z0-9])", normalized)) * 2.5
    return score


def row_string_from_cells(row: list[str]) -> str:
    return "".join(cell or " " for cell in row[:COLS]).rstrip()


def empty_grid() -> list[list[str]]:
    return [["" for _ in range(COLS)] for _ in range(ROWS)]


def empty_confidence_grid() -> list[list[float]]:
    return [[0.0 for _ in range(COLS)] for _ in range(ROWS)]


def clamp_data_col(col: int) -> int:
    return max(FIRST_DATA_COL, min(LAST_DATA_COL, col))


def nearest_data_col(x: float, cell_w: float) -> int:
    return clamp_data_col(int(round(x / cell_w)))


def normalize_grid_guards(grid: list[list[str]]) -> list[list[str]]:
    normalized = [[cell for cell in row[:COLS]] + [""] * max(0, COLS - len(row)) for row in grid[:ROWS]]
    for row in normalized:
        row[0] = ""
        row[COLS - 1] = ""
    return normalized


def calibrate_axis(centers: list[float], nominal_pitch: float, count: int) -> tuple[float, float]:
    if len(centers) < 6:
        return 0.0, nominal_pitch
    values = np.asarray(centers, dtype=np.float32)
    best = (float("inf"), 0.0, nominal_pitch)
    for scale in np.linspace(0.985, 1.015, 13):
        pitch = nominal_pitch * float(scale)
        indexes = np.clip(np.round(values / pitch - 0.5), 0, count - 1)
        offsets = values - (indexes + 0.5) * pitch
        origin = max(-nominal_pitch * 0.16, min(nominal_pitch * 0.16, float(np.median(offsets))))
        indexes = np.clip(np.round((values - origin) / pitch - 0.5), 0, count - 1)
        residuals = np.abs(values - (origin + (indexes + 0.5) * pitch))
        score = float(np.median(residuals)) + abs(origin) * 0.35 + abs(scale - 1.0) * nominal_pitch
        if score < best[0]:
            best = (score, origin, pitch)
    return best[1], best[2]


def calibrate_grid(
    char_boxes: list[dict[str, Any]],
    screen_size: tuple[int, int],
) -> dict[str, float]:
    screen_w, screen_h = screen_size
    x_centers = [
        float(box["left"]) + float(box["width"]) * 0.5
        for box in char_boxes
        if 2 <= float(box.get("width") or 0) <= screen_w / COLS * 2.1
    ]
    y_centers = [
        float(box["top"]) + float(box["height"]) * 0.5
        for box in char_boxes
        if 2 <= float(box.get("height") or 0) <= screen_h / ROWS * 1.4
    ]
    nominal_cell_w = screen_w / COLS
    nominal_cell_h = screen_h / ROWS
    origin_x, _ = calibrate_axis(x_centers, nominal_cell_w, COLS)
    origin_y, _ = calibrate_axis(y_centers, nominal_cell_h, ROWS)
    # The UI grid is exactly 40 by 13. OCR must use the identical pitch; even a
    # 1% fitted scale error accumulates into a full-column shift at the right edge.
    return {
        "origin_x": origin_x,
        "origin_y": origin_y,
        "cell_w": nominal_cell_w,
        "cell_h": nominal_cell_h,
    }


def place_char(grid: list[list[str]], char_box: dict[str, Any], geometry: dict[str, float]) -> bool:
    text = clean_ocr_text(str(char_box["text"]))[:1]
    if not text:
        return False

    cell_w = geometry["cell_w"]
    cell_h = geometry["cell_h"]
    box_w = float(char_box.get("width") or 0)
    box_h = float(char_box.get("height") or 0)
    if box_w > cell_w * 2.15 or box_h > cell_h * 1.35 or box_w < 2 or box_h < 2:
        return False

    center_x = float(char_box["left"]) + box_w * 0.5
    center_y = float(char_box["top"]) + box_h * 0.5
    row = max(0, min(ROWS - 1, int((center_y - geometry["origin_y"]) / cell_h)))
    col = clamp_data_col(int((center_x - geometry["origin_x"]) / cell_w))

    if grid[row][col] and grid[row][col] != text:
        for offset in (-1, 1, -2, 2):
            next_col = col + offset
            if FIRST_DATA_COL <= next_col <= LAST_DATA_COL and not grid[row][next_col]:
                col = next_col
                break
    if not grid[row][col]:
        grid[row][col] = text
        return True
    return False


def build_character_grid(
    char_boxes: list[dict[str, Any]],
    geometry: dict[str, float],
) -> tuple[list[list[str]], list[list[float]]]:
    grid = empty_grid()
    confidence = empty_confidence_grid()
    cell_w = geometry["cell_w"]
    cell_h = geometry["cell_h"]
    for box in sorted(char_boxes, key=lambda item: (float(item["top"]), float(item["left"]))):
        text = clean_ocr_text(str(box.get("text", "")))[:1]
        box_w = float(box.get("width") or 0)
        box_h = float(box.get("height") or 0)
        if not text or box_w < 2 or box_h < 2 or box_w > cell_w * 2.15 or box_h > cell_h * 1.4:
            continue
        center_x = float(box["left"]) + box_w * 0.5
        center_y = float(box["top"]) + box_h * 0.5
        row = max(0, min(ROWS - 1, int((center_y - geometry["origin_y"]) / cell_h)))
        col = clamp_data_col(int((center_x - geometry["origin_x"]) / cell_w))
        score = min(0.95, 0.50 + min(0.35, box_h / max(1.0, cell_h) * 0.35))
        if not grid[row][col] or score > confidence[row][col]:
            grid[row][col] = text
            confidence[row][col] = score
    return normalize_grid_guards(grid), confidence


def place_word(
    grid: list[list[str]],
    scores: list[list[float]],
    confidence: list[list[float]],
    word: dict[str, Any],
    geometry: dict[str, float],
) -> None:
    text = normalize_mcdu_phrase(str(word["text"]))
    if not text:
        return

    cell_w = geometry["cell_w"]
    cell_h = geometry["cell_h"]
    row = max(
        0,
        min(
            ROWS - 1,
            int((word["top"] + word["height"] * 0.5 - geometry["origin_y"]) / cell_h),
        ),
    )
    start_col = clamp_data_col(int(round((float(word["left"]) - geometry["origin_x"]) / cell_w)))

    compact = text.replace(" ", "")
    if not compact:
        return

    quality = word_quality(word)
    box_cols = max(1, int(round(float(word.get("width") or 1) / cell_w)))
    sequence = text if " " in text and len(text) <= box_cols + 2 else compact
    use_projected_spacing = " " not in sequence and box_cols > len(compact) + 1
    max_len = LAST_DATA_COL - start_col + 1

    for index, char in enumerate(sequence[:max_len]):
        if char.isspace():
            continue
        if use_projected_spacing:
            center_x = float(word["left"]) + float(word["width"]) * ((index + 0.5) / len(compact))
            col = clamp_data_col(int((center_x - geometry["origin_x"]) / cell_w))
        else:
            col = start_col + index
        if not grid[row][col] or quality > scores[row][col] + 10:
            grid[row][col] = char
            scores[row][col] = quality
            confidence[row][col] = max(0.20, min(0.98, float(word.get("conf") or 0.0) / 100.0))


def merge_ocr_grids(
    word_grid: list[list[str]],
    word_confidence: list[list[float]],
    char_grid: list[list[str]],
    char_confidence: list[list[float]],
) -> tuple[list[list[str]], list[list[float]]]:
    merged = empty_grid()
    confidence = empty_confidence_grid()
    for row in range(ROWS):
        for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
            word = word_grid[row][col]
            char = char_grid[row][col]
            if char and word == char:
                merged[row][col] = char
                confidence[row][col] = min(1.0, max(char_confidence[row][col], word_confidence[row][col]) + 0.08)
            elif char and not word:
                nearby = range(max(FIRST_DATA_COL, col - 2), min(LAST_DATA_COL, col + 2) + 1)
                same_nearby = any(word_grid[row][nearby_col] == char for nearby_col in nearby)
                left_context = any(word_grid[row][nearby_col] for nearby_col in range(max(FIRST_DATA_COL, col - 2), col))
                right_context = any(
                    word_grid[row][nearby_col]
                    for nearby_col in range(col + 1, min(LAST_DATA_COL, col + 2) + 1)
                )
                adjacent = "".join(
                    word_grid[row][nearby_col] or " "
                    for nearby_col in range(max(FIRST_DATA_COL, col - 1), min(LAST_DATA_COL, col + 1) + 1)
                )
                degree_context = char == "°" and any(value.isdigit() for value in adjacent)
                if degree_context or (left_context and right_context and not same_nearby):
                    merged[row][col] = char
                    confidence[row][col] = min(char_confidence[row][col], 0.72)
            elif word:
                merged[row][col] = word
                confidence[row][col] = word_confidence[row][col] if not char else min(word_confidence[row][col], 0.58)
    return normalize_grid_guards(merged), confidence


def apply_char_box_spacing(
    grid: list[list[str]],
    boxes: list[dict[str, Any]],
    geometry: dict[str, float],
) -> list[list[str]]:
    """B4: Restore spaces hidden by word-mode packing using per-character box positions.

    Tesseract word mode returns one bounding box per word token, losing internal
    spaces. When char boxes show a gap ≥ 1.7 × cell_w between the centres of boxes
    assigned to adjacent columns, a blank cell should separate them. Cells right of
    the gap are shifted one column rightward — but only when LAST_DATA_COL is empty
    so no content is lost.
    """
    if not boxes:
        return grid

    cell_w = geometry["cell_w"]
    cell_h = geometry["cell_h"]
    origin_x = geometry["origin_x"]
    origin_y = geometry["origin_y"]

    updated = [row[:] for row in grid]

    row_entries: dict[int, list[tuple[int, float]]] = {}
    for box in boxes:
        box_w = float(box.get("width") or 0)
        box_h = float(box.get("height") or 0)
        if box_w < 2 or box_h < 2 or box_w > cell_w * 2.15 or box_h > cell_h * 1.4:
            continue
        cx = float(box["left"]) + box_w * 0.5
        cy = float(box["top"]) + box_h * 0.5
        row_idx = max(0, min(ROWS - 1, int((cy - origin_y) / cell_h)))
        col_idx = clamp_data_col(int((cx - origin_x) / cell_w))
        row_entries.setdefault(row_idx, []).append((col_idx, cx))

    for row_idx, entries in row_entries.items():
        entries.sort(key=lambda t: t[1])
        for i in range(len(entries) - 1):
            col_a, cx_a = entries[i]
            col_b, cx_b = entries[i + 1]
            if col_b != col_a + 1:
                continue
            if not updated[row_idx][col_a] or not updated[row_idx][col_b]:
                continue
            if cx_b - cx_a <= cell_w * 1.7:
                continue
            if updated[row_idx][LAST_DATA_COL]:
                continue
            for c in range(LAST_DATA_COL, col_b, -1):
                updated[row_idx][c] = updated[row_idx][c - 1]
            updated[row_idx][col_b] = ""

    return updated


def fuse_engine_grids(
    primary_grid: list[list[str]],
    primary_confidence: list[list[float]],
    secondary_grid: list[list[str]],
    secondary_confidence: list[list[float]],
) -> tuple[list[list[str]], list[list[float]], dict[str, int]]:
    fused = normalize_grid_guards(primary_grid)
    confidence = [row[:] for row in primary_confidence]
    summary = {"rowsSelected": 0, "agreements": 0, "blanksFilled": 0, "disagreements": 0}
    for row in range(ROWS):
        primary_runs: dict[int, int] = {}
        col = FIRST_DATA_COL
        while col <= LAST_DATA_COL:
            if not primary_grid[row][col]:
                col += 1
                continue
            start = col
            while col <= LAST_DATA_COL and primary_grid[row][col]:
                col += 1
            for run_col in range(start, col):
                primary_runs[run_col] = col - start

        secondary_chars = [
            col
            for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1)
            if secondary_grid[row][col]
        ]
        secondary_average = (
            sum(secondary_confidence[row][col] for col in secondary_chars) / len(secondary_chars)
            if secondary_chars
            else 0.0
        )
        primary_text = row_string_from_cells(primary_grid[row])
        secondary_text = row_string_from_cells(secondary_grid[row])
        primary_is_separator = primary_text.count("-") >= 8
        use_secondary_row = (
            len(secondary_chars) >= 2
            and secondary_average >= 0.80
            and not (primary_is_separator and secondary_text.count("-") < 4)
        )
        if use_secondary_row:
            fused[row] = secondary_grid[row][:]
            confidence[row] = secondary_confidence[row][:]
            summary["rowsSelected"] += 1

        for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
            primary = primary_grid[row][col]
            secondary = secondary_grid[row][col]
            if primary and secondary and primary == secondary:
                summary["agreements"] += 1
                confidence[row][col] = min(
                    1.0,
                    max(confidence[row][col], secondary_confidence[row][col]) + 0.10,
                )
            elif (
                not use_secondary_row
                and not primary
                and secondary
                and secondary_confidence[row][col] >= 0.50
            ):
                summary["blanksFilled"] += 1
                fused[row][col] = secondary
                confidence[row][col] = min(0.78, secondary_confidence[row][col])
            elif primary and secondary and primary != secondary:
                summary["disagreements"] += 1
                confidence[row][col] = min(confidence[row][col], 0.45)
            elif (
                use_secondary_row
                and primary
                and not secondary
                and primary_confidence[row][col] >= 0.60
                and (primary in {"-", "°", "<", ">"} or primary_runs.get(col, 0) >= 2)
            ):
                nearby = range(max(FIRST_DATA_COL, col - 2), min(LAST_DATA_COL, col + 2) + 1)
                if not any(secondary_grid[row][nearby_col] for nearby_col in nearby):
                    fused[row][col] = primary
                    confidence[row][col] = min(0.72, primary_confidence[row][col])
                    summary["blanksFilled"] += 1
    return normalize_grid_guards(fused), confidence, summary


def fuse_atlas_grid(
    grid: list[list[str]],
    confidence: list[list[float]],
    atlas_grid: list[list[str]],
    atlas_confidence: list[list[float]],
) -> tuple[list[list[str]], list[list[float]]]:
    """Fill empty cells from atlas results (B1, fill-only).

    The atlas may only write to a cell that is currently empty in *grid*.  Any
    cell that already carries text — regardless of its confidence — is left
    unchanged.  Atlas confidence is capped at _ATLAS_MAX_CONFIDENCE (0.55) so
    it never outranks a later Tesseract correction.
    """
    out = [row[:] for row in grid]
    out_conf = [row[:] for row in confidence]
    for row in range(ROWS):
        for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
            if not atlas_grid[row][col]:
                continue
            if not out[row][col]:
                out[row][col] = atlas_grid[row][col]
                out_conf[row][col] = min(atlas_confidence[row][col], _ATLAS_MAX_CONFIDENCE)
    return out, out_conf


def connected_components(mask: np.ndarray) -> list[dict[str, Any]]:
    height, width = mask.shape
    seen = np.zeros(mask.shape, dtype=bool)
    components: list[dict[str, Any]] = []
    for start_y in range(height):
        for start_x in range(width):
            if seen[start_y, start_x] or not mask[start_y, start_x]:
                continue
            stack = [(start_x, start_y)]
            seen[start_y, start_x] = True
            xs: list[int] = []
            ys: list[int] = []
            while stack:
                x, y = stack.pop()
                xs.append(x)
                ys.append(y)
                for next_x, next_y in ((x + 1, y), (x - 1, y), (x, y + 1), (x, y - 1)):
                    if (
                        0 <= next_x < width
                        and 0 <= next_y < height
                        and not seen[next_y, next_x]
                        and mask[next_y, next_x]
                    ):
                        seen[next_y, next_x] = True
                        stack.append((next_x, next_y))
            count = len(xs)
            if count < 250:
                continue
            x1 = min(xs)
            x2 = max(xs) + 1
            y1 = min(ys)
            y2 = max(ys) + 1
            components.append(
                {
                    "count": count,
                    "x1": x1,
                    "y1": y1,
                    "x2": x2,
                    "y2": y2,
                    "fill": count / max(1, (x2 - x1) * (y2 - y1)),
                    "xs": xs,
                    "ys": ys,
                }
            )
    return components


def robust_line_fit(independent: np.ndarray, dependent: np.ndarray) -> tuple[float, float] | None:
    if independent.size < 12 or dependent.size != independent.size:
        return None
    keep = np.ones(independent.size, dtype=bool)
    slope = 0.0
    intercept = float(np.median(dependent))
    for _ in range(3):
        if int(np.sum(keep)) < 8:
            return None
        slope, intercept = np.polyfit(independent[keep], dependent[keep], 1)
        residuals = np.abs(dependent - (slope * independent + intercept))
        cutoff = max(1.5, float(np.percentile(residuals[keep], 72)))
        keep = residuals <= cutoff
    return float(slope), float(intercept)


def intersect_edge_lines(
    vertical: tuple[float, float],
    horizontal: tuple[float, float],
) -> tuple[float, float] | None:
    # vertical: x = a*y + b; horizontal: y = c*x + d
    a, b = vertical
    c, d = horizontal
    denominator = 1.0 - a * c
    if abs(denominator) < 1e-6:
        return None
    x = (a * d + b) / denominator
    y = c * x + d
    return float(x), float(y)


def regularize_quadrilateral(points: list[tuple[float, float]]) -> list[tuple[float, float]]:
    if len(points) != 4:
        return points
    tl, tr, br, bl = points
    top_width = math.dist(tl, tr)
    bottom_width = math.dist(bl, br)
    reference_width = max(1.0, top_width, bottom_width)
    left_dx = bl[0] - tl[0]
    right_dx = br[0] - tr[0]
    width_ratio = max(top_width, bottom_width) / max(1.0, min(top_width, bottom_width))
    vertical_divergence = abs(left_dx - right_dx) / reference_width

    # Dark-pixel components can be split by a selected magenta field. This most
    # often drags only the lower-left fit into the display while the other three
    # edges remain accurate. Reconstruct that single outlier by parallelogram
    # closure; ordinary phone perspective is retained unless both checks fail.
    left_is_outlier = abs(left_dx) > abs(right_dx) * 3.0 + reference_width * 0.03
    if width_ratio > 1.12 and vertical_divergence > 0.12 and left_is_outlier:
        predicted_bl = (tl[0] + br[0] - tr[0], tl[1] + br[1] - tr[1])
        return [tl, tr, br, predicted_bl]
    return points


def polygon_area(points: list[tuple[float, float]]) -> float:
    return 0.5 * abs(
        sum(
            points[index][0] * points[(index + 1) % len(points)][1]
            - points[(index + 1) % len(points)][0] * points[index][1]
            for index in range(len(points))
        )
    )


def line_intersection(
    first: tuple[float, float, float],
    second: tuple[float, float, float],
) -> tuple[float, float] | None:
    a1, b1, c1 = first
    a2, b2, c2 = second
    denominator = a1 * b2 - a2 * b1
    if abs(denominator) < 1e-6:
        return None
    return (
        (b1 * c2 - b2 * c1) / denominator,
        (c1 * a2 - c2 * a1) / denominator,
    )


def segment_line(segment: np.ndarray) -> tuple[float, float, float]:
    x1, y1, x2, y2 = (float(value) for value in segment)
    a = y1 - y2
    b = x2 - x1
    length = max(1e-6, math.hypot(a, b))
    return a / length, b / length, (x1 * y2 - x2 * y1) / length


def refine_display_corners(
    image: Image.Image,
    corners: list[dict[str, float]],
) -> list[dict[str, float]]:
    """Snap a rough dark-region box to the four physical display border lines."""
    try:
        import cv2
    except (ImportError, OSError):
        return corners

    if len(corners) != 4:
        return corners
    points = np.asarray([[corner["x"], corner["y"]] for corner in corners], dtype=np.float32)
    gray = cv2.cvtColor(np.asarray(image), cv2.COLOR_RGB2GRAY)
    gray = cv2.GaussianBlur(gray, (5, 5), 0)
    edges = cv2.Canny(gray, 35, 115)

    widths = [float(np.linalg.norm(points[1] - points[0])), float(np.linalg.norm(points[2] - points[3]))]
    heights = [float(np.linalg.norm(points[3] - points[0])), float(np.linalg.norm(points[2] - points[1]))]
    min_edge = max(30.0, min(widths + heights))
    lines = cv2.HoughLinesP(
        edges,
        1,
        np.pi / 720,
        threshold=max(35, int(min_edge * 0.13)),
        minLineLength=max(25, int(min_edge * 0.32)),
        maxLineGap=max(8, int(min_edge * 0.05)),
    )
    if lines is None:
        return corners

    # Bottom and left use reversed endpoints so every expected side follows the
    # clockwise display outline. Direction itself is ignored during matching.
    expected = [
        (points[0], points[1], sum(heights) * 0.5),
        (points[1], points[2], sum(widths) * 0.5),
        (points[3], points[2], sum(heights) * 0.5),
        (points[0], points[3], sum(widths) * 0.5),
    ]
    selected: list[tuple[float, float, float]] = []
    for start, end, opposite_size in expected:
        vector = end - start
        expected_length = float(np.linalg.norm(vector))
        if expected_length < 20:
            return corners
        unit = vector / expected_length
        normal = np.asarray([-unit[1], unit[0]], dtype=np.float32)
        tolerance = max(8.0, opposite_size * 0.085)
        best: tuple[float, np.ndarray] | None = None
        for raw_segment in lines[:, 0]:
            segment = raw_segment.astype(np.float32)
            segment_vector = segment[2:4] - segment[0:2]
            segment_length = float(np.linalg.norm(segment_vector))
            if segment_length < expected_length * 0.24:
                continue
            segment_unit = segment_vector / max(1e-6, segment_length)
            angle = math.degrees(math.acos(min(1.0, abs(float(np.dot(unit, segment_unit))))))
            if angle > 13.0:
                continue
            midpoint = (segment[0:2] + segment[2:4]) * 0.5
            along = float(np.dot(midpoint - start, unit))
            if along < -expected_length * 0.16 or along > expected_length * 1.16:
                continue
            distance = abs(float(np.dot(midpoint - start, normal)))
            if distance > tolerance:
                continue
            coverage = min(1.0, segment_length / expected_length)
            score = distance / tolerance + angle / 13.0 - coverage * 0.62
            if best is None or score < best[0]:
                best = (score, segment)
        if best is None:
            return corners
        selected.append(segment_line(best[1]))

    top, right, bottom, left = selected
    intersections = [
        line_intersection(left, top),
        line_intersection(top, right),
        line_intersection(right, bottom),
        line_intersection(bottom, left),
    ]
    if any(point is None for point in intersections):
        return corners
    refined = [point for point in intersections if point is not None]
    rough = [(float(point[0]), float(point[1])) for point in points]
    rough_area = max(1.0, polygon_area(rough))
    refined_area = polygon_area(refined)
    max_shift = max(max(widths), max(heights)) * 0.13
    if not (rough_area * 0.70 <= refined_area <= rough_area * 1.30):
        return corners
    if any(math.dist(old, new) > max_shift for old, new in zip(rough, refined)):
        return corners
    if any(x < 0 or y < 0 or x >= image.width or y >= image.height for x, y in refined):
        return corners
    return [{"x": float(x), "y": float(y)} for x, y in refined]


def _candidate_corners(
    component: dict[str, Any], scale: float, small_w: int, small_h: int
) -> list[dict[str, float]]:
    """Apply a tiny inset to a component's bounding box, then fit corner lines."""
    pad_x = max(1, round((component["x2"] - component["x1"]) * 0.003))
    pad_y = max(1, round((component["y2"] - component["y1"]) * 0.003))
    padded = dict(component)
    padded.update(
        {
            "x1": max(0, component["x1"] + pad_x),
            "y1": max(0, component["y1"] + pad_y),
            "x2": min(small_w, component["x2"] - pad_x),
            "y2": min(small_h, component["y2"] - pad_y),
        }
    )
    return component_corners(padded, scale)


def component_corners(component: dict[str, Any], scale: float) -> list[dict[str, float]]:
    x1 = float(component["x1"])
    y1 = float(component["y1"])
    x2 = float(component["x2"])
    y2 = float(component["y2"])
    fallback = [(x1, y1), (x2, y1), (x2, y2), (x1, y2)]
    xs = np.asarray(component.get("xs") or [], dtype=np.float32)
    ys = np.asarray(component.get("ys") or [], dtype=np.float32)
    inside = (xs >= x1) & (xs <= x2) & (ys >= y1) & (ys <= y2)
    xs = xs[inside]
    ys = ys[inside]
    if xs.size < 100 or ys.size < 100:
        points = fallback
    else:
        width = max(1.0, x2 - x1)
        height = max(1.0, y2 - y1)
        unique_y = np.unique(ys.astype(np.int32))
        left_x: list[float] = []
        right_x: list[float] = []
        edge_y: list[float] = []
        for y in unique_y:
            row_x = xs[ys.astype(np.int32) == y]
            if row_x.size < 4:
                continue
            edge_y.append(float(y))
            left_x.append(float(np.percentile(row_x, 2)))
            right_x.append(float(np.percentile(row_x, 98)))

        unique_x = np.unique(xs.astype(np.int32))
        top_y: list[float] = []
        bottom_y: list[float] = []
        edge_x: list[float] = []
        for x in unique_x:
            col_y = ys[xs.astype(np.int32) == x]
            if col_y.size < 4:
                continue
            edge_x.append(float(x))
            top_y.append(float(np.percentile(col_y, 2)))
            bottom_y.append(float(np.percentile(col_y, 98)))

        edge_y_arr = np.asarray(edge_y, dtype=np.float32)
        edge_x_arr = np.asarray(edge_x, dtype=np.float32)
        left_line = robust_line_fit(edge_y_arr, np.asarray(left_x, dtype=np.float32))
        right_line = robust_line_fit(edge_y_arr, np.asarray(right_x, dtype=np.float32))
        top_line = robust_line_fit(edge_x_arr, np.asarray(top_y, dtype=np.float32))
        bottom_line = robust_line_fit(edge_x_arr, np.asarray(bottom_y, dtype=np.float32))
        intersections = (
            [
                intersect_edge_lines(left_line, top_line),
                intersect_edge_lines(right_line, top_line),
                intersect_edge_lines(right_line, bottom_line),
                intersect_edge_lines(left_line, bottom_line),
            ]
            if left_line and right_line and top_line and bottom_line
            else []
        )
        if not intersections or any(point is None for point in intersections):
            points = fallback
        else:
            points = [point for point in intersections if point is not None]
            margin_x = width * 0.12
            margin_y = height * 0.12
            corner_shift = max(
                max(abs(px - ex) / width, abs(py - ey) / height)
                for (px, py), (ex, ey) in zip(points, fallback)
            )
            if corner_shift > 0.20 or any(
                px < x1 - margin_x or px > x2 + margin_x or py < y1 - margin_y or py > y2 + margin_y
                for px, py in points
            ):
                points = fallback
            else:
                polygon_area = 0.5 * abs(
                    sum(
                        points[index][0] * points[(index + 1) % 4][1]
                        - points[(index + 1) % 4][0] * points[index][1]
                        for index in range(4)
                    )
                )
                if polygon_area < width * height * 0.62:
                    points = fallback

    points = regularize_quadrilateral(points)

    inv = 1 / scale
    return [{"x": x * inv, "y": y * inv} for x, y in points]


def detect_display(payload: dict[str, Any]) -> dict[str, Any]:
    image = load_image(str(payload["image"]))

    # Orientation probe: run before detection so detection sees the corrected image.
    # Skipped when the caller already rotated the image manually (skipOrientationProbe).
    rotation = 0
    if not bool(payload.get("skipOrientationProbe", False)):
        try:
            rotation = probe_orientation(image)
        except Exception:  # noqa: BLE001 - probe failure must not block detection
            rotation = 0
        if rotation != 0:
            image = image.rotate(rotation, expand=True)

    scale = min(1.0, 900 / max(image.size))
    small = image.resize((max(1, round(image.width * scale)), max(1, round(image.height * scale))))
    arr = np.asarray(small).astype(np.float32)
    gray = arr[:, :, 0] * 0.299 + arr[:, :, 1] * 0.587 + arr[:, :, 2] * 0.114
    mask = gray < 58
    components = connected_components(mask)
    if not components:
        raise ValueError("Could not find a dark MCDU display region.")

    image_area = small.width * small.height
    scored: list[tuple[float, dict[str, Any], bool]] = []
    for component in components:
        width = component["x2"] - component["x1"]
        height = component["y2"] - component["y1"]
        area = width * height
        aspect = width / max(1, height)
        # Hard filters: reject absurdly small or non-screen-shaped blobs
        if area < image_area * 0.05 or aspect < 0.7 or aspect > 4.5:
            continue
        touches_edge = (
            component["x1"] <= 2
            or component["y1"] <= 2
            or component["x2"] >= small.width - 2
            or component["y2"] >= small.height - 2
        )
        base_score = component["count"] * min(1.0, component["fill"] * 1.4)
        # Soft bonus for aspect ratios in the MCDU landscape range (≈1.4–1.9)
        aspect_factor = max(0.55, 1.0 - max(0.0, abs(aspect - 1.65) - 0.5) * 0.18)
        # Completeness: edge-touching blobs are likely cut-off neighbour displays
        completeness = 0.30 if touches_edge else 1.0
        scored.append((base_score * aspect_factor * completeness, component, touches_edge))

    if not scored:
        raise ValueError("Could not isolate the black display. Drag the four corners manually.")

    scored.sort(key=lambda item: -item[0])
    best_score = scored[0][0]

    # Keep candidates with at least 12 % of the top score, capped at 6
    candidates_raw = [
        entry for entry in scored[:6] if entry[0] >= best_score * 0.12
    ]

    # Build candidate descriptors: component_corners for all, Hough refinement only for best
    candidates: list[dict[str, Any]] = []
    perspective_refined = False
    for i, (score, component, touches_edge) in enumerate(candidates_raw):
        corners = _candidate_corners(component, scale, small.width, small.height)
        if i == 0:
            small_corners = [{"x": c["x"] * scale, "y": c["y"] * scale} for c in corners]
            refined_small = refine_display_corners(small, small_corners)
            perspective_refined = any(
                edge_length(before, after) > 1.5
                for before, after in zip(small_corners, refined_small)
            )
            corners = [{"x": c["x"] / scale, "y": c["y"] / scale} for c in refined_small]
        candidates.append(
            {
                "corners": corners,
                "confidence": round(float(component["fill"]), 3),
                "score": round(score, 1),
                "boundingBox": {
                    "x": round(component["x1"] / scale),
                    "y": round(component["y1"] / scale),
                    "width": round((component["x2"] - component["x1"]) / scale),
                    "height": round((component["y2"] - component["y1"]) / scale),
                },
                "touchesEdge": touches_edge,
            }
        )

    best_corners = candidates[0]["corners"]
    display_width = (
        edge_length(best_corners[0], best_corners[1]) + edge_length(best_corners[3], best_corners[2])
    ) * 0.5
    display_height = (
        edge_length(best_corners[0], best_corners[3]) + edge_length(best_corners[1], best_corners[2])
    ) * 0.5
    return {
        "corners": best_corners,
        "confidence": candidates[0]["confidence"],
        "perspectiveRefined": perspective_refined,
        "displaySize": {"width": round(display_width), "height": round(display_height)},
        "candidates": candidates,
        "bestIndex": 0,
        "rotation": rotation,
    }


def flatten_display(payload: dict[str, Any]) -> dict[str, Any]:
    image = load_image(str(payload["image"]))
    corners = payload.get("corners")
    if not isinstance(corners, list) or len(corners) != 4:
        raise ValueError("Four screen corners are required.")

    warped, grid_origin = warp_screen_with_alignment(image, corners)
    preview_id = f"{uuid.uuid4().hex}.png"
    preview_path = EXPORTS / preview_id
    warped.save(preview_path)
    return {
        "previewUrl": f"/data/exports/{preview_id}",
        "width": warped.width,
        "height": warped.height,
        "gridAlignment": {
            "x": grid_origin[0] / warped.width,
            "y": grid_origin[1] / warped.height,
        },
    }


def load_corrections() -> dict[str, str]:
    ensure_dirs()
    try:
        data = json.loads(CORRECTIONS.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        data = {}
    return {str(key): str(value) for key, value in data.items()}


def save_corrections(corrections: dict[str, str]) -> None:
    ensure_dirs()
    CORRECTIONS.write_text(json.dumps(corrections, indent=2, sort_keys=True), encoding="utf-8")


def grid_row_text(grid: list[Any], row: int) -> str:
    values = grid[row] if row < len(grid) and isinstance(grid[row], list) else []
    return "".join(str(values[col])[:1] if col < len(values) and values[col] else " " for col in range(COLS))


def normalize_requirement_value(value: str) -> str:
    return re.sub(r"\s+", " ", clean_ocr_text(value)).strip().upper()


def requirement_matches(requirement: dict[str, Any], observed: str) -> bool:
    expected = str(requirement.get("expected", ""))
    ignore_case = bool(requirement.get("ignoreCase", True))
    ignore_spaces = bool(requirement.get("ignoreSpaces", False))
    check_type = str(requirement.get("type", "exact"))

    def normalized(value: str) -> str:
        value = clean_ocr_text(value)
        if ignore_spaces:
            value = re.sub(r"\s+", "", value)
        if ignore_case:
            value = value.upper()
        return value

    actual = normalized(observed)
    target = normalized(expected)
    if check_type == "blank":
        return not actual.strip()
    if check_type == "contains":
        return target in actual
    if check_type == "not_contains":
        return target not in actual
    if check_type == "fill":
        width = int(requirement["end"]) - int(requirement["start"]) + 1
        return actual == normalized(expected * width)
    return actual == target


def focused_grid_read(
    warped: Image.Image,
    row: int,
    start: int,
    end: int,
    expected: str,
) -> str:
    cell_w = warped.width / COLS
    cell_h = warped.height / ROWS
    x1 = max(0, int(start * cell_w))
    x2 = min(warped.width, int((end + 1) * cell_w))
    y1 = max(0, int((row - 1) * cell_h))
    y2 = min(warped.height, int(row * cell_h))
    crop = warped.crop((x1, y1, x2, y2))
    scale = 4
    enlarged = crop.resize((max(1, crop.width * scale), max(1, crop.height * scale)), Image.Resampling.LANCZOS)
    processed = preprocess_for_ocr(enlarged)
    width = end - start + 1
    focused = [" " for _ in range(width)]
    focused_cell_w = enlarged.width / max(1, width)

    try:
        words = run_tesseract_tsv(processed)
    except RuntimeError:
        words = []

    for word in words:
        text = normalize_mcdu_phrase(str(word.get("text", "")))
        if not text:
            continue
        compact = text.replace(" ", "")
        if not compact:
            continue
        word_left = float(word.get("left") or 0)
        word_width = float(word.get("width") or 1)
        start_index = max(0, min(width - 1, int(round(word_left / focused_cell_w))))
        span_cols = max(1, int(round(word_width / focused_cell_w)))
        sequence = text if " " in text and len(text) <= span_cols + 2 else compact
        use_projected_spacing = " " not in sequence and span_cols > len(compact) + 1
        for index, char in enumerate(sequence):
            if char.isspace():
                continue
            if use_projected_spacing:
                center_x = word_left + word_width * ((index + 0.5) / max(1, len(compact)))
                target = max(0, min(width - 1, int(center_x / focused_cell_w)))
            else:
                target = start_index + index
            if 0 <= target < width and focused[target] == " ":
                focused[target] = char

    if not "".join(focused).strip():
        for box in run_tesseract_boxes(processed):
            box_w = float(box.get("width") or 0)
            box_h = float(box.get("height") or 0)
            if box_w < 2 or box_h < 2 or box_w > focused_cell_w * 2.1:
                continue
            center_x = float(box["left"]) + box_w * 0.5
            index = max(0, min(width - 1, int(center_x / focused_cell_w)))
            char = clean_ocr_text(str(box.get("text", "")))[:1]
            if char and focused[index] == " ":
                focused[index] = char

    if expected == "-":
        gray = np.asarray(ImageOps.grayscale(crop)).astype(np.float32)
        for index in range(width):
            cx1 = int(index * crop.width / width)
            cx2 = int((index + 1) * crop.width / width)
            cy1 = int(crop.height * 0.36)
            cy2 = int(crop.height * 0.68)
            cell = gray[cy1:cy2, cx1:cx2]
            if cell.size == 0:
                continue
            bright = cell > 170
            if float(np.max(np.mean(bright, axis=1))) >= 0.30:
                focused[index] = "-"

    cell_reading = "".join(focused)
    raw_reading = " ".join(clean_ocr_text(str(word["text"])) for word in words).strip()
    if cell_reading.strip():
        return cell_reading
    return raw_reading[:width].ljust(width)


def grid_from_payload(value: Any) -> list[list[str]]:
    if not isinstance(value, list) or len(value) != ROWS:
        raise ValueError("A 13-row grid is required.")
    grid = empty_grid()
    for row_index, row in enumerate(value[:ROWS]):
        if not isinstance(row, list):
            continue
        for col_index, cell in enumerate(row[:COLS]):
            grid[row_index][col_index] = clean_ocr_text(str(cell))[:1] if cell else ""
    return normalize_grid_guards(grid)


def whole_grid_focused_recheck(
    warped: Image.Image,
    grid: list[list[str]],
    mode: str,
) -> tuple[list[list[str]], dict[str, int]]:
    mode = mode if mode in {"conservative", "balanced", "aggressive"} else "conservative"
    updated = normalize_grid_guards([[cell for cell in row] for row in grid])
    summary = {
        "rowsChecked": 0,
        "cellsFilled": 0,
        "cellsReplaced": 0,
        "dashesRecovered": 0,
        "conflictsKept": 0,
        "rowsSkipped": 0,
    }
    allowed = set("ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789<>/().°-+ ")

    for row in range(1, ROWS + 1):
        summary["rowsChecked"] += 1
        focused = focused_grid_read(warped, row, FIRST_DATA_COL, LAST_DATA_COL, "")
        focused = focused[: LAST_DATA_COL - FIRST_DATA_COL + 1].ljust(LAST_DATA_COL - FIRST_DATA_COL + 1)
        current_row = updated[row - 1]
        current_snapshot = current_row[:]
        current_compact = "".join(cell or "" for cell in current_row[FIRST_DATA_COL : LAST_DATA_COL + 1])
        focused_compact = "".join(char for char in focused if char and not char.isspace())
        current_text = row_string_from_cells(current_row)
        focused_text = f" {focused} "
        focused_score = mcdu_row_score(focused_text)
        current_score = mcdu_row_score(current_text)
        if (
            mode in {"balanced", "aggressive"}
            and focused_compact
            and focused_score >= 35.0
            and len(focused_compact) >= max(3, int(len(current_compact) * 0.55))
            and focused_score >= current_score + 10.0
        ):
            for offset, char in enumerate(focused):
                col = FIRST_DATA_COL + offset
                updated[row - 1][col] = "" if char.isspace() else clean_ocr_text(char)[:1]
            summary["cellsReplaced"] += sum(
                1
                for offset, char in enumerate(focused)
                if clean_ocr_text(char).strip()
                and current_snapshot[FIRST_DATA_COL + offset] != clean_ocr_text(char)[:1]
            )
            continue
        if mode != "aggressive" and (current_score >= 20.0 or focused_score < 25.0):
            continue
        if (
            mode != "aggressive"
            and current_compact
            and focused_compact
            and difflib.SequenceMatcher(None, current_compact, focused_compact).ratio() < 0.65
        ):
            summary["rowsSkipped"] += 1
            continue

        for offset, char in enumerate(focused):
            col = FIRST_DATA_COL + offset
            char = clean_ocr_text(char)[:1]
            if not char or char.isspace() or char not in allowed:
                continue
            current = updated[row - 1][col]
            if not current:
                has_left_context = col > FIRST_DATA_COL and bool(updated[row - 1][col - 1])
                has_right_context = col < LAST_DATA_COL and bool(updated[row - 1][col + 1])
                if mode != "aggressive" and not (has_left_context and has_right_context):
                    continue
                updated[row - 1][col] = char
                if char == "-":
                    summary["dashesRecovered"] += 1
                else:
                    summary["cellsFilled"] += 1
                continue
            if current == char:
                continue
            if mode == "aggressive":
                updated[row - 1][col] = char
                summary["cellsReplaced"] += 1
            else:
                summary["conflictsKept"] += 1

    updated = recover_dash_lines(updated, warped) if mode in {"balanced", "aggressive"} else updated
    updated = disambiguate_o_zero(updated)
    updated = validate_field_formats(updated)  # D2: same snapping as main analyze path
    updated = snap_title_row(updated)  # A5: title vocabulary snap
    updated = apply_corrections(updated)
    return normalize_grid_guards(updated), summary


def review_requirements(payload: dict[str, Any]) -> dict[str, Any]:
    requirements = payload.get("requirements")
    grid = payload.get("grid")
    if not isinstance(requirements, list) or not requirements:
        raise ValueError("Add at least one requirement before reviewing.")
    if not isinstance(grid, list) or len(grid) != ROWS:
        raise ValueError("Analyze or enter a 13-row grid before reviewing requirements.")

    warped: Image.Image | None = None
    image_data = str(payload.get("image", ""))
    corners = payload.get("corners")
    if image_data and isinstance(corners, list) and len(corners) == 4:
        warped = warp_screen(load_image(image_data), corners)

    results: list[dict[str, Any]] = []
    for index, raw_requirement in enumerate(requirements, 1):
        if not isinstance(raw_requirement, dict):
            continue
        try:
            row = int(raw_requirement["row"])
            start = int(raw_requirement["start"])
            end = int(raw_requirement["end"])
        except (KeyError, TypeError, ValueError):
            row = start = end = 0
        check_type = str(raw_requirement.get("type", "exact"))
        expected_raw = str(raw_requirement.get("expected", ""))
        valid_type = check_type in {"exact", "contains", "fill", "blank", "not_contains"}
        valid_expected = check_type == "blank" or bool(expected_raw)
        valid_fill = check_type != "fill" or len(expected_raw) == 1
        if not (1 <= row <= ROWS and 1 <= start <= end <= 38 and valid_type and valid_expected and valid_fill):
            results.append(
                {
                    "line": index,
                    "requirement": "Invalid requirement",
                    "status": "NEEDS REVIEW",
                    "expected": expected_raw,
                    "observed": "",
                    "rechecked": "",
                    "location": "",
                    "detail": "Check the row, column range, check type, and expected text.",
                }
            )
            continue

        requirement = dict(raw_requirement)
        width = end - start + 1
        expected_display = (
            expected_raw * width
            if check_type == "fill"
            else "Blank"
            if check_type == "blank"
            else expected_raw
        )
        observed = grid_row_text(grid, row - 1)[start : end + 1]
        passed = requirement_matches(requirement, observed)
        rechecked = ""
        detail = "Matched the extracted grid." if passed else "Initial grid reading did not match."

        if not passed and warped is not None:
            rechecked = focused_grid_read(warped, row, start, end, expected_raw)
            recheck_passed = requirement_matches(requirement, rechecked)
            detail = (
                "Focused OCR matched, but the initial grid did not. Confirm this result."
                if recheck_passed
                else "Focused OCR recheck also did not match."
            )
        else:
            recheck_passed = False

        type_label = {
            "exact": "Exact text",
            "contains": "Contains text",
            "fill": "Fill range",
            "blank": "Blank range",
            "not_contains": "Must not contain",
        }[check_type]
        results.append(
            {
                "line": index,
                "requirement": type_label,
                "status": "PASS" if passed else "NEEDS REVIEW" if recheck_passed else "FAIL",
                "expected": expected_display,
                "observed": observed,
                "rechecked": rechecked,
                "location": f"Row {row}, columns {start}-{end}",
                "detail": detail,
            }
        )

    summary = {
        "total": len(results),
        "passed": sum(result["status"] == "PASS" for result in results),
        "failed": sum(result["status"] == "FAIL" for result in results),
        "needsReview": sum(result["status"] == "NEEDS REVIEW" for result in results),
    }
    return {"results": results, "summary": summary}


def compact_row_text(value: str) -> str:
    return re.sub(r"\s+", "", clean_ocr_text(value)).upper()


def fixed_row_text(value: str) -> str:
    return value[:COLS].ljust(COLS)


def correction_key(row: int, original: str) -> str:
    return f"row:{row + 1}|{fixed_row_text(original)}"


def apply_corrections(grid: list[list[str]]) -> list[list[str]]:
    corrections = load_corrections()
    updated = [[cell for cell in row] for row in grid]
    for row_index, row in enumerate(updated):
        row_text = fixed_row_text("".join(cell or " " for cell in row))
        exact_key = correction_key(row_index, row_text)
        if exact_key in corrections:
            updated[row_index] = list(fixed_row_text(corrections[exact_key]))
    return normalize_grid_guards(updated)


def apply_templates(grid: list[list[str]], warped: Image.Image) -> list[list[str]]:
    templates = load_templates()
    if not templates:
        return grid
    updated = [[cell for cell in row] for row in grid]
    gray = np.asarray(ImageOps.grayscale(warped)).astype(np.float32)
    screen_w, screen_h = warped.size
    cell_w = screen_w / COLS
    cell_h = screen_h / ROWS
    for row in range(ROWS):
        for col in range(COLS):
            x1 = int(col * cell_w)
            x2 = int((col + 1) * cell_w)
            y1 = int(row * cell_h)
            y2 = int((row + 1) * cell_h)
            crop = gray[y1:y2, x1:x2]
            current = updated[row][col]
            if crop.size == 0 or (float(crop.max()) < 105 and current != "-"):
                continue
            if col < FIRST_DATA_COL or col > LAST_DATA_COL:
                continue
            classified = classify_from_templates(cell_feature(warped, row, col), templates)
            if not classified:
                continue
            char, distance = classified
            if char == BLANK_TEMPLATE_KEY:
                if current == "-" and distance <= 0.10:
                    updated[row][col] = ""
                continue
            if not current and char in {"-", "<", ">", ".", "/"}:
                continue
            if not current or (current in {"O", "0"} and char in {"O", "0"} and distance <= 0.10):
                updated[row][col] = char
    return updated


def normalize_o_zero_token(token: str) -> str:
    if not token or not any(char in token for char in "O0"):
        return token

    letters_only = re.sub(r"[^A-Z0]", "", token.upper())
    as_word = letters_only.replace("0", "O")
    if as_word in MCDU_VOCABULARY:
        return token.replace("0", "O")

    if re.fullmatch(r"FL[O0-9]{2,3}", token):
        return "FL" + token[2:].replace("O", "0")

    result = list(token)
    digit_count = sum(value.isdigit() for value in token)
    letter_count = sum(value.isalpha() for value in token)
    for index, char in enumerate(result):
        if char not in {"O", "0"}:
            continue
        left = result[index - 1] if index > 0 else ""
        right = result[index + 1] if index + 1 < len(result) else ""
        numeric_context = (
            left.isdigit()
            or right.isdigit()
            or (digit_count > 0 and any(mark in token for mark in (".", "/", "°")))
            or digit_count > letter_count
        )
        result[index] = "0" if numeric_context else "O"
    return "".join(result)


def disambiguate_o_zero(grid: list[list[str]]) -> list[list[str]]:
    updated = [[cell for cell in row] for row in grid]
    for row in range(ROWS):
        col = FIRST_DATA_COL
        while col <= LAST_DATA_COL:
            if not updated[row][col]:
                col += 1
                continue
            start = col
            chars: list[str] = []
            while col <= LAST_DATA_COL and updated[row][col]:
                chars.append(updated[row][col])
                col += 1
            normalized = normalize_o_zero_token("".join(chars))
            for offset, char in enumerate(normalized):
                updated[row][start + offset] = char
    return normalize_grid_guards(updated)


def grid_character_count(grid: list[list[str]]) -> int:
    return sum(1 for row in grid for cell in row[FIRST_DATA_COL : LAST_DATA_COL + 1] if cell)


def verify_grid_with_char_boxes(
    word_grid: list[list[str]],
    char_boxes: list[dict[str, Any]],
    geometry: dict[str, float],
) -> list[list[str]]:
    verified = [[cell for cell in row] for row in word_grid]
    char_grid = empty_grid()
    for box in char_boxes:
        place_char(char_grid, box, geometry)

    for row in range(ROWS):
        for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
            char_value = char_grid[row][col]
            if not char_value or verified[row][col]:
                continue
            nearby_same = any(
                verified[row][nearby] == char_value
                for nearby in range(max(FIRST_DATA_COL, col - 2), min(LAST_DATA_COL, col + 2) + 1)
            )
            if nearby_same:
                verified[row][col] = char_value
    return normalize_grid_guards(verified)


def recover_dash_lines(grid: list[list[str]], warped: Image.Image) -> list[list[str]]:
    updated = [[cell for cell in row] for row in grid]
    gray = np.asarray(ImageOps.grayscale(warped)).astype(np.float32)
    screen_w, screen_h = warped.size
    cell_w = screen_w / COLS
    cell_h = screen_h / ROWS

    # Robust separator pass. The MCDU dashed divider appears as a single thin
    # bright horizontal line spanning most of the width with little other ink.
    # OCR otherwise reads it as random characters, so detect it straight from the
    # image and force the whole data span to dashes, overriding any garbage. The
    # full-width line (cover ≈ 0.8) is well separated from text rows (cover ≤ 0.5)
    # and from short entry-box edges, so the threshold is safe.
    data_x1 = max(0, int(FIRST_DATA_COL * cell_w))
    data_x2 = min(screen_w, int((LAST_DATA_COL + 1) * cell_w))
    for row in range(ROWS):
        y1 = max(0, int(row * cell_h))
        y2 = min(screen_h, int((row + 1) * cell_h))
        band = gray[y1:y2, data_x1:data_x2]
        if band.size == 0 or band.shape[0] == 0:
            continue
        bright = band > 160
        line_cover = float(bright.mean(axis=1).max())
        total_ink = float(bright.mean())
        if line_cover >= 0.6 and total_ink <= 0.22:
            for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
                updated[row][col] = "-"

    for row in range(ROWS):
        existing_dash_cols = [
            col for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1) if updated[row][col] == "-"
        ]
        if len(existing_dash_cols) >= 12:
            start = min(existing_dash_cols)
            end = max(existing_dash_cols)
            non_dash = sum(
                1 for col in range(start, end + 1) if updated[row][col] and updated[row][col] != "-"
            )
            if non_dash <= 3:
                for col in range(start, end + 1):
                    updated[row][col] = "-"

        y1 = int((row + 0.25) * cell_h)
        y2 = int((row + 0.80) * cell_h)
        if y2 <= y1:
            continue
        strip = gray[max(0, y1) : min(screen_h, y2), :]
        if strip.size == 0:
            continue

        active_cols: list[int] = []
        relaxed_dash_cols: list[int] = []
        for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
            x1 = int(col * cell_w + cell_w * 0.10)
            x2 = int((col + 1) * cell_w - cell_w * 0.10)
            cell = strip[:, max(0, x1) : min(screen_w, x2)]
            if cell.size == 0:
                continue
            bright = cell > 178
            bright_ratio = float(np.mean(bright))
            row_stroke = float(np.max(np.mean(bright, axis=1)))
            if 0.006 <= bright_ratio <= 0.16 and row_stroke >= 0.34:
                active_cols.append(col)
            if 0.003 <= bright_ratio <= 0.30 and row_stroke >= 0.45:
                relaxed_dash_cols.append(col)

        run_start: int | None = None
        runs: list[tuple[int, int]] = []
        previous = -99
        for col in active_cols:
            if run_start is None or col != previous + 1:
                if run_start is not None and previous - run_start + 1 >= 9:
                    runs.append((run_start, previous))
                run_start = col
            previous = col
        if run_start is not None and previous - run_start + 1 >= 9:
            runs.append((run_start, previous))

        for start, end in runs:
            # Avoid turning a text-heavy row into dashes. Separator rows normally have
            # long mostly empty spans with only a few labels at the edges.
            has_text_left = any(
                str(updated[row][col]).isalpha() for col in range(FIRST_DATA_COL, start)
            )
            has_text_right = any(
                str(updated[row][col]).isalpha() for col in range(end + 1, LAST_DATA_COL + 1)
            )
            if has_text_left and has_text_right:
                continue
            occupied = sum(1 for col in range(start, end + 1) if updated[row][col])
            dash_count = sum(1 for col in range(start, end + 1) if updated[row][col] == "-")
            non_dash_count = occupied - dash_count
            if end - start + 1 >= 12 and dash_count >= 6 and non_dash_count <= 3:
                for col in range(start, end + 1):
                    updated[row][col] = "-"
                continue
            if occupied > max(3, (end - start + 1) // 4):
                continue
            for col in range(start, end + 1):
                if not updated[row][col]:
                    updated[row][col] = "-"

        current_dash_count = sum(
            updated[row][col] == "-" for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1)
        )
        if current_dash_count >= 12 and len(relaxed_dash_cols) >= 12:
            relaxed_start = min(relaxed_dash_cols)
            relaxed_end = max(relaxed_dash_cols)
            for col in range(relaxed_start, relaxed_end + 1):
                updated[row][col] = "-"

    # A row may become dash-dominant only after the image pass fills missed cells.
    # Normalize once more so isolated OCR punctuation cannot survive inside it.
    for row in range(ROWS):
        dash_cols = [col for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1) if updated[row][col] == "-"]
        if len(dash_cols) < 12:
            continue
        start = min(dash_cols)
        end = max(dash_cols)
        non_dash = sum(1 for col in range(start, end + 1) if updated[row][col] and updated[row][col] != "-")
        if non_dash <= 3:
            for col in range(start, end + 1):
                updated[row][col] = "-"
    return normalize_grid_guards(updated)


# ---------------------------------------------------------------------------
# #22 — Per-field format validator
# #23 — Color-semantics extraction
# #25 — Blur / quality gate
# ---------------------------------------------------------------------------


def compute_blur_score(warped: Image.Image) -> float:
    """Variance-of-Laplacian blur metric for the warped screen (#25).

    Returns a non-negative float — higher means sharper.  Values below
    BLUR_THRESHOLD (~80) indicate that the image may be too blurry for reliable
    OCR.  Returns 0.0 when OpenCV is not available.
    """
    try:
        import cv2
    except ImportError:
        return 0.0
    gray = np.asarray(ImageOps.grayscale(warped), dtype=np.uint8)
    laplacian = cv2.Laplacian(gray, cv2.CV_64F)
    return float(laplacian.var())


def fuse_grids(grids: list[list[list[str]]]) -> tuple[list[list[str]], dict[str, Any]]:
    """Majority-vote fusion across 2–3 pre-analyzed grids (#24).

    For each cell, collects the non-empty votes from all grids and picks
    the most frequent character.  If only one grid is non-empty that value
    wins (fill-in).  When all grids are empty the cell stays empty.

    Returns ``(fused_grid, summary)`` where summary counts agreed / filled /
    conflicted / empty cells across the whole grid.
    """
    if len(grids) < 2:
        raise ValueError("At least two grids are required for fusion.")
    for g in grids:
        if len(g) != ROWS or any(len(row) != COLS for row in g):
            raise ValueError(f"All grids must be {ROWS}×{COLS}.")

    fused: list[list[str]] = empty_grid()
    agreed = filled = conflicted = empty = 0

    for row in range(ROWS):
        for col in range(COLS):
            votes = [g[row][col] for g in grids if g[row][col]]
            if not votes:
                empty += 1
                continue
            # A4: deterministic tie-breaking — preserve first-occurrence order so
            # that equal-count characters resolve to the one from the earliest grid.
            unique_votes = list(dict.fromkeys(votes))
            best = max(unique_votes, key=votes.count)
            fused[row][col] = best
            if len(set(votes)) == 1:
                if len(votes) == len(grids):
                    agreed += 1
                else:
                    filled += 1
            else:
                conflicted += 1

    summary = {
        "grids": len(grids),
        "agreed": agreed,
        "filled": filled,
        "conflicted": conflicted,
        "empty": empty,
    }
    return fused, summary


def _snap_fl_token(text: str) -> str | None:
    """Snap to FL\\d{2,3} via one-character substitution; None if not applicable.

    Covers first-char confusions (E/P/T → F) and second-char confusions
    (I/1/J → L).  Only same-length corrections are returned so column
    alignment is always preserved.
    """
    if re.fullmatch(r"FL[0-9]{2,3}", text):
        return None
    if re.fullmatch(r"F[A-Z1][0-9]{2,3}", text):   # wrong second char
        return "FL" + text[2:]
    if re.fullmatch(r"[A-EG-Z]L[0-9]{2,3}", text):  # wrong first char
        return "FL" + text[2:]
    return None


def _snap_decimal_token(text: str) -> str | None:
    """Snap to .\\d{3} via one-character substitution; None if not applicable.

    Fixes comma-instead-of-dot, the most common OCR error for Mach speeds.
    """
    if re.fullmatch(r"\.[0-9]{3}", text):
        return None
    if re.fullmatch(r",[0-9]{3}", text):
        return "." + text[1:]
    return None


def validate_field_formats(grid: list[list[str]]) -> list[list[str]]:
    """Apply one-edit format corrections guided by row-label context (#22).

    Two passes:
    1. Standalone — snap FL-like tokens whose first or second character was
       substituted, and snap comma-for-dot decimal tokens.  Safe to apply
       without label context because no valid MCDU token looks like
       ``[A-Z]L\\d{2,3}`` except an FL altitude.
    2. Label-driven — when the previous row contains an altitude label
       (CRZ / OPT / MAX / RECMD) or a speed label (ECON / LRC / SEL / RTA
       SPD), re-run the corresponding validator on all tokens in the current
       row.  This covers the common 777-9 layout where labels and values
       appear on consecutive rows.
    """
    updated = [[cell for cell in row] for row in grid]

    def get_tokens(r: int) -> list[tuple[int, str]]:
        result: list[tuple[int, str]] = []
        col = FIRST_DATA_COL
        while col <= LAST_DATA_COL:
            if not updated[r][col]:
                col += 1
                continue
            start = col
            chars: list[str] = []
            while col <= LAST_DATA_COL and updated[r][col]:
                chars.append(updated[r][col])
                col += 1
            result.append((start, "".join(chars)))
        return result

    def apply_snap(r: int, start: int, old: str, new: str) -> None:
        if not new or new == old or len(new) != len(old):
            return
        for i, ch in enumerate(new):
            updated[r][start + i] = ch

    _alt_re = re.compile(r"\b(?:CRZ|OPT|MAX|RECMD)\b")
    _spd_re = re.compile(
        r"\b(?:ECON|LRC|SEL|RTA)\s*SPD\b|\bSPD\s*(?:ECON|LRC|SEL|RTA)\b"
    )
    # Unambiguous cruise-Mach speed labels (SEL SPD can be a knots value).
    _mach_re = re.compile(r"\b(?:ECON|LRC|RTA)\s*SPD\b|\bSPD\s*(?:ECON|LRC|RTA)\b")

    # Pass 1: standalone token-level snapping (no label context needed)
    for row in range(ROWS):
        for start, token in get_tokens(row):
            snapped = _snap_fl_token(token) or _snap_decimal_token(token)
            if snapped:
                apply_snap(row, start, token, snapped)

    # Pass 2: label-driven validation (label on row N → value on row N+1)
    for row in range(1, ROWS):
        label_text = " ".join(t for _, t in get_tokens(row - 1))
        if _alt_re.search(label_text):
            for start, token in get_tokens(row):
                snapped = _snap_fl_token(token)
                if snapped:
                    apply_snap(row, start, token, snapped)
        elif _spd_re.search(label_text):
            for start, token in get_tokens(row):
                snapped = _snap_decimal_token(token)
                if snapped:
                    apply_snap(row, start, token, snapped)
                elif (
                    re.fullmatch(r"\d{3}", token)
                    and _mach_re.search(label_text)
                    and start > FIRST_DATA_COL
                    and not updated[row][start - 1]
                ):
                    # Cruise Mach speed is shown as ".733"; OCR routinely drops the
                    # faint leading dot. Restore it only for the unambiguous Mach
                    # labels (ECON/LRC/RTA) — SEL SPD can be a knots value (e.g. 173)
                    # that must not gain a dot.
                    updated[row][start - 1] = "."

    return updated


def detect_title_chip_text(
    warped: Image.Image, geometry: dict[str, float]
) -> list[tuple[int, str]]:
    """A5 Part 1: OCR the inverse-video ACT/MOD chip in the title row.

    The chip is a light-background rectangle with dark text. The global image
    inversion used by the main OCR pipeline maps it to a dark background, causing
    Tesseract to drop or mangle the text. This function crops the chip directly from
    the warped image (correct polarity) and runs a dedicated Tesseract pass.
    Returns (col_idx, char) pairs for row 0.
    """
    try:
        import cv2
    except ImportError:
        return []

    cell_w = geometry["cell_w"]
    cell_h = geometry["cell_h"]
    origin_x = geometry["origin_x"]
    origin_y = geometry["origin_y"]
    img_w, img_h = warped.size

    row_y0 = max(0, int(origin_y))
    row_y1 = min(img_h, int(origin_y + cell_h))
    if row_y1 <= row_y0:
        return []

    rgb = np.asarray(warped.convert("RGB"), dtype=np.uint8)
    strip = rgb[row_y0:row_y1, :, :]
    hsv_strip = cv2.cvtColor(strip, cv2.COLOR_RGB2HSV)

    # Bright (V > 160) and low-saturation (S < 50) = white/gray chip background
    bright_mask = cv2.inRange(
        hsv_strip,
        np.array([0, 0, 160], dtype=np.uint8),
        np.array([179, 50, 255], dtype=np.uint8),
    )
    min_chip_area = max(int(cell_w * cell_h * 1.5), 200)
    num_labels, _labels, stats, _ = cv2.connectedComponentsWithStats(bright_mask, connectivity=8)

    results: list[tuple[int, str]] = []
    for i in range(1, num_labels):
        area = int(stats[i, cv2.CC_STAT_AREA])
        if area < min_chip_area:
            continue
        bx = int(stats[i, cv2.CC_STAT_LEFT])
        bw = int(stats[i, cv2.CC_STAT_WIDTH])
        bh = int(stats[i, cv2.CC_STAT_HEIGHT])
        if bw < int(cell_w * 1.5) or bh < int(cell_h * 0.4):
            continue
        if area / max(1, bw * bh) < 0.55:
            continue
        chip_rgb = strip[:, bx : bx + bw, :]
        if chip_rgb.size == 0:
            continue
        chip_img = Image.fromarray(chip_rgb, mode="RGB")
        padded = ImageOps.expand(chip_img, border=4, fill=(255, 255, 255))
        for word in run_tesseract_tsv(padded):
            text = clean_ocr_text(str(word.get("text", "")))
            if not text:
                continue
            wx = int(word.get("left", 0)) - 4 + bx
            ww = int(word.get("width", 0))
            cx = wx + ww * 0.5
            col = clamp_data_col(int((cx - origin_x) / cell_w))
            for j, ch in enumerate(text):
                results.append((clamp_data_col(col + j), ch))
    return results


def snap_title_row(grid: list[list[str]]) -> list[list[str]]:
    """A5 Part 2: Fuzzy-snap row 0 text to the nearest MCDU_TITLE_VOCABULARY entry.

    Reads the non-empty cells of row 0, compares the joined text against the closed
    vocabulary with difflib, and rewrites those cells only when there is an
    unambiguous close match (ratio ≥ 0.70 over best, margin ≥ 0.10 over second-best).
    Correct titles produce ratio == 1.0 and are left byte-identical; garbage that
    matches nothing closely is untouched.
    """
    updated = [row[:] for row in grid]
    row0 = updated[0]

    non_empty = [c for c in range(FIRST_DATA_COL, LAST_DATA_COL + 1) if row0[c]]
    if not non_empty:
        return updated

    start_col = non_empty[0]
    end_col = non_empty[-1]
    raw_text = "".join(row0[c] or " " for c in range(start_col, end_col + 1)).strip()
    if not raw_text:
        return updated

    best_r = 0.0
    best_entry: str | None = None
    second_r = 0.0
    for entry in MCDU_TITLE_VOCABULARY:
        r = difflib.SequenceMatcher(None, raw_text, entry).ratio()
        if r > best_r:
            second_r = best_r
            best_r = r
            best_entry = entry
        elif r > second_r:
            second_r = r

    if best_entry is None or best_r < 0.70 or best_r - second_r < 0.10:
        return updated

    for c in range(start_col, min(start_col + len(best_entry), LAST_DATA_COL + 1)):
        ch = best_entry[c - start_col]
        row0[c] = ch if ch != " " else ""
    for c in range(start_col + len(best_entry), end_col + 1):
        row0[c] = ""

    return updated


def _classify_cell_color(hsv_crop: np.ndarray) -> str:
    """Return the dominant text color in a cell's HSV crop (#23).

    Looks only at bright pixels (V > 120) which represent text, not the dark
    background.  Returns one of "white", "magenta", "cyan", "amber", or "".
    """
    v_mask = hsv_crop[:, :, 2] > 120
    if not np.any(v_mask):
        return ""
    saturations = hsv_crop[:, :, 1][v_mask].astype(np.float32)
    if float(np.mean(saturations)) < 45:
        return "white"
    hues = hsv_crop[:, :, 0][v_mask].astype(np.float32)
    median_hue = float(np.median(hues))
    if median_hue >= 130:
        return "magenta"
    if 80 <= median_hue < 130:
        return "cyan"
    if 15 <= median_hue < 80:
        return "amber"
    return "white"


def extract_color_semantics(warped: Image.Image, grid: list[list[str]]) -> list[list[str]]:
    """Return a color-label grid for each non-empty cell in the OCR result (#23).

    Each non-empty cell is labelled "white", "magenta", "cyan", "amber", or "".
    Magenta = modified/active value; cyan = boxed/selected; amber = caution.
    Returns an all-empty grid when OpenCV is not available.
    """
    try:
        import cv2
    except ImportError:
        return [[""] * COLS for _ in range(ROWS)]

    rgb = np.asarray(warped.convert("RGB"), dtype=np.uint8)
    hsv = cv2.cvtColor(rgb, cv2.COLOR_RGB2HSV)
    img_h, img_w = rgb.shape[:2]
    cell_w = img_w / COLS
    cell_h = img_h / ROWS

    color_grid: list[list[str]] = [[""] * COLS for _ in range(ROWS)]
    for row in range(ROWS):
        for col in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
            if not grid[row][col]:
                continue
            x1 = int(col * cell_w)
            x2 = int((col + 1) * cell_w)
            y1 = int(row * cell_h)
            y2 = int((row + 1) * cell_h)
            crop = hsv[max(0, y1) : min(img_h, y2), max(0, x1) : min(img_w, x2)]
            if crop.size == 0:
                continue
            color_grid[row][col] = _classify_cell_color(crop)
    return color_grid


def remember_templates(payload: dict[str, Any]) -> dict[str, Any]:
    image = load_image(str(payload["image"]))
    corners = payload.get("corners")
    grid = payload.get("grid")
    source_grid = payload.get("sourceGrid")
    if not isinstance(corners, list) or len(corners) != 4:
        raise ValueError("Four screen corners are required for template learning.")
    if not isinstance(grid, list) or len(grid) != ROWS:
        raise ValueError("A corrected 13-row grid is required for template learning.")
    if not isinstance(source_grid, list) or len(source_grid) != ROWS:
        raise ValueError("The original OCR grid is required for selective template learning.")

    warped = warp_screen(image, corners)
    templates = load_templates()
    learned = 0
    for row_index, row in enumerate(grid):
        if not isinstance(row, list):
            continue
        for col_index, value in enumerate(row[:COLS]):
            if col_index < FIRST_DATA_COL or col_index > LAST_DATA_COL:
                continue
            char = clean_ocr_text(str(value))[:1]
            original = ""
            if isinstance(source_grid[row_index], list) and col_index < len(source_grid[row_index]):
                original = clean_ocr_text(str(source_grid[row_index][col_index]))[:1]
            if char == original:
                continue
            if not char or char.isspace():
                if original == "-":
                    templates.setdefault(BLANK_TEMPLATE_KEY, []).append(
                        cell_feature(warped, row_index, col_index)
                    )
                    learned += 1
                continue
            templates.setdefault(char, []).append(cell_feature(warped, row_index, col_index))
            learned += 1
    save_templates(templates)
    return {"learned": learned, "characters": len(templates)}


def deduplicate_adjacent_rows(
    word_grid: list[list[str]],
    word_confidence: list[list[float]],
    word_candidates: list[tuple[str, list[list[str]], list[list[float]]]],
    row_sources: list[str],
) -> None:
    """Replace rows that duplicate their neighbour above with the best non-duplicate candidate (A1).

    Operates in-place.  For each row whose text is identical to the row above,
    re-ranks ``word_candidates`` by plain score and picks the highest-scoring
    candidate whose text differs from the row above.  If no such candidate
    exists the row is blanked (empty strings), which is safer than showing
    a spurious label repeat.
    """
    for row in range(1, ROWS):
        above_text = row_string_from_cells(word_grid[row - 1]).strip()
        this_text = row_string_from_cells(word_grid[row]).strip()
        if not (above_text and this_text and above_text == this_text):
            continue

        def _score(candidate: tuple[str, list[list[str]], list[list[float]]], _r: int = row) -> float:
            _, cg, cc = candidate
            t = row_string_from_cells(cg[_r])
            return mcdu_row_score(t) + sum(cc[_r][FIRST_DATA_COL : LAST_DATA_COL + 1]) * 2.0

        ranked = sorted(word_candidates, key=_score, reverse=True)
        replaced = False
        for cname, cg, cc in ranked:
            if row_string_from_cells(cg[row]).strip() != above_text:
                word_grid[row] = cg[row][:]
                word_confidence[row] = cc[row][:]
                row_sources[row] = cname
                replaced = True
                break
        if not replaced:
            word_grid[row] = [""] * COLS
            word_confidence[row] = [0.0] * COLS


def analyze(payload: dict[str, Any]) -> dict[str, Any]:
    image = load_image(str(payload["image"]))
    corners = payload.get("corners")
    if not isinstance(corners, list) or len(corners) != 4:
        raise ValueError("Four screen corners are required.")

    warped = warp_screen(image, corners)
    warped = clean_warped_image(warped)  # #14/#15/#17: cursor / glare / outlines
    blur_score = compute_blur_score(warped)  # #25: quality gate
    screen_size = warped.size

    # C1: run all variant TSV passes in parallel
    _variants = preprocessing_variants(warped)
    _t0 = time.perf_counter()
    with ThreadPoolExecutor(max_workers=min(len(_variants), 8)) as _pool:
        _all_words = list(_pool.map(lambda vi: run_tesseract_tsv(vi[1]), _variants))
    _t_tsv = time.perf_counter() - _t0
    ocr_candidates: list[tuple[float, str, Image.Image, list[dict[str, Any]]]] = [
        (ocr_words_score(vwords, vimg.height), vname, vimg, vwords)
        for (vname, vimg), vwords in zip(_variants, _all_words)
    ]
    _, preprocessing_name, processed, words = max(ocr_candidates, key=lambda item: item[0])
    # warp_screen already phase-aligns the image to the visible 40 x 13 lattice.
    # Re-fitting an OCR offset here would make extraction disagree with the grid
    # the user sees and edits in the browser.
    geometry = calibrate_grid([], screen_size)

    word_candidates: list[tuple[str, list[list[str]], list[list[float]]]] = []
    for _, variant_name, _, variant_words in ocr_candidates:
        candidate_grid = empty_grid()
        candidate_scores = [[0.0 for _ in range(COLS)] for _ in range(ROWS)]
        candidate_confidence = empty_confidence_grid()
        for word in variant_words:
            place_word(candidate_grid, candidate_scores, candidate_confidence, word, geometry)
        word_candidates.append((variant_name, candidate_grid, candidate_confidence))

    # #16: OCR coloured message boxes (blue/cyan/amber) without global inversion
    box_words = ocr_message_boxes(warped, geometry)
    if box_words:
        box_grid = empty_grid()
        box_scores: list[list[float]] = [[0.0] * COLS for _ in range(ROWS)]
        box_confidence = empty_confidence_grid()
        for word in box_words:
            place_word(box_grid, box_scores, box_confidence, word, geometry)
        word_candidates.append(("message_box", box_grid, box_confidence))

    # C3: skip per-row strip if existing word candidates are already high-confidence
    _draft_conf = [[0.0] * COLS for _ in range(ROWS)]
    for _, _cg, _cc in word_candidates:
        for _r in range(ROWS):
            for _c in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
                if _cg[_r][_c] and _cc[_r][_c] > _draft_conf[_r][_c]:
                    _draft_conf[_r][_c] = _cc[_r][_c]
    _filled_conf = [
        _draft_conf[_r][_c]
        for _r in range(ROWS)
        for _c in range(FIRST_DATA_COL, LAST_DATA_COL + 1)
        if _draft_conf[_r][_c] > 0
    ]
    _skip_strip = bool(_filled_conf) and sum(_filled_conf) / len(_filled_conf) >= _STRIP_SKIP_CONFIDENCE

    # C1: run box pass and (optionally) per-row strip concurrently
    _t1 = time.perf_counter()
    with ThreadPoolExecutor(max_workers=2) as _pool:
        _box_fut = _pool.submit(run_tesseract_boxes, processed)
        _strip_fut = None if _skip_strip else _pool.submit(per_row_strip_ocr, warped, geometry)
        boxes = _box_fut.result()
        strip_grid, strip_conf = (
            _strip_fut.result() if _strip_fut is not None else (empty_grid(), empty_confidence_grid())
        )
    _t_box_strip = time.perf_counter() - _t1

    # #18: Per-row strip OCR — add to candidates (empty placeholder if C3 skip)
    word_candidates.append(("per_row", strip_grid, strip_conf))
    print(
        f"[ocr-timing] variants={len(_variants)}"
        f" tsv={_t_tsv:.2f}s"
        f" box+strip={_t_box_strip:.2f}s"
        f" strip={'skipped(C3)' if _skip_strip else 'run'}"
    )

    hybrid_requested = bool(payload.get("hybridOcr", True))
    paddle_words: list[dict[str, Any]] = []
    paddle_error = ""
    paddle_grid = empty_grid()
    paddle_confidence = empty_confidence_grid()
    if hybrid_requested:
        paddle_words, paddle_error = run_paddle_ocr(warped)
        if paddle_words:
            paddle_scores = [[0.0 for _ in range(COLS)] for _ in range(ROWS)]
            for word in paddle_words:
                place_word(paddle_grid, paddle_scores, paddle_confidence, word, geometry)

    word_grid = empty_grid()
    word_confidence = empty_confidence_grid()
    row_sources: list[str] = []
    for row in range(ROWS):
        def row_candidate_score(candidate: tuple[str, list[list[str]], list[list[float]]], _row: int = row) -> float:
            _, candidate_grid, candidate_confidence = candidate
            text = row_string_from_cells(candidate_grid[_row])
            confidence_score = sum(candidate_confidence[_row][FIRST_DATA_COL : LAST_DATA_COL + 1])
            character_count = sum(bool(cell) for cell in candidate_grid[_row][FIRST_DATA_COL : LAST_DATA_COL + 1])
            score = mcdu_row_score(text) + confidence_score * 2.0 + character_count * 0.5
            # A1: penalize repeating the row just selected above, and bonus digit
            # content that logically follows a vocabulary label row.
            if _row > 0:
                above_text = row_string_from_cells(word_grid[_row - 1]).strip()
                this_text = text.strip()
                if above_text and this_text and above_text == this_text:
                    score -= 30.0
                if above_text and any(w in above_text for w in MCDU_VOCABULARY if len(w) >= 3):
                    if re.search(r"\d", text):
                        score += 8.0
            return score

        source_name, source_grid, source_confidence = max(word_candidates, key=row_candidate_score)
        word_grid[row] = source_grid[row][:]
        word_confidence[row] = source_confidence[row][:]
        row_sources.append(source_name)

    # A1: post-selection safety net — replace any remaining adjacent duplicates
    deduplicate_adjacent_rows(word_grid, word_confidence, word_candidates, row_sources)

    fusion_summary = None
    if paddle_words:
        word_grid, word_confidence, fusion_summary = fuse_engine_grids(
            word_grid,
            word_confidence,
            paddle_grid,
            paddle_confidence,
        )
    char_grid, char_confidence = build_character_grid(boxes, geometry)
    grid, confidence_grid = merge_ocr_grids(word_grid, word_confidence, char_grid, char_confidence)
    grid = apply_char_box_spacing(grid, boxes, geometry)  # B4

    # B1: atlas glyph classifier (fill-only: only writes to cells empty after
    # Tesseract/PaddleOCR merge).  Disabled: cross-font accuracy caps at ~75 %
    # with available system fonts, below the 85 % production threshold.
    if ATLAS_ENABLED:
        _atlas = _get_glyph_atlas()
        if _atlas:
            _atlas_grid = empty_grid()
            _atlas_conf = empty_confidence_grid()
            for _ar in range(ROWS):
                for _ac in range(FIRST_DATA_COL, LAST_DATA_COL + 1):
                    _res = classify_cell_atlas(warped, _ar, _ac, _atlas)
                    if _res is not None:
                        _char, _score = _res
                        _atlas_grid[_ar][_ac] = _char
                        _atlas_conf[_ar][_ac] = _score
            grid, confidence_grid = fuse_atlas_grid(
                grid, confidence_grid, _atlas_grid, _atlas_conf
            )

    grid = recover_dash_lines(grid, warped)
    grid = apply_templates(grid, warped)
    grid = disambiguate_o_zero(grid)
    grid = validate_field_formats(grid)  # #22: per-field format snapping
    # A5: re-OCR inverse-video title chip then snap title to closed vocabulary
    for _chip_col, _chip_ch in detect_title_chip_text(warped, geometry):
        grid[0][_chip_col] = _chip_ch
    grid = snap_title_row(grid)
    corrected_grid = apply_corrections(grid)
    verification_summary = None
    verification = payload.get("verification")
    if isinstance(verification, dict) and bool(verification.get("enabled")):
        corrected_grid, verification_summary = whole_grid_focused_recheck(
            warped,
            corrected_grid,
            str(verification.get("mode", "conservative")),
        )

    color_grid = extract_color_semantics(warped, corrected_grid)  # #23: color labels

    preview_id = f"{uuid.uuid4().hex}.png"
    preview_path = EXPORTS / preview_id
    warped.save(preview_path)

    return {
        "grid": corrected_grid,
        "words": words,
        "boxes": boxes,
        "calibration": geometry,
        "confidenceGrid": confidence_grid,
        "colorGrid": color_grid,
        "preprocessing": preprocessing_name,
        "rowPreprocessing": row_sources,
        "ocrEngines": {
            "requested": ["tesseract", "paddle"] if hybrid_requested else ["tesseract"],
            "used": ["tesseract", "paddle"] if paddle_words else ["tesseract"],
            "paddleWords": len(paddle_words),
            "paddleError": paddle_error,
            "fusion": fusion_summary,
        },
        "previewUrl": f"/data/exports/{preview_id}",
        "verification": verification_summary,
        "blurScore": round(blur_score, 1),
        "blurry": blur_score > 0 and blur_score < BLUR_THRESHOLD,
        "blurWarning": BLUR_THRESHOLD <= blur_score < BLUR_WARNING_THRESHOLD,  # D3: marginal band
    }


def refine_grid(payload: dict[str, Any]) -> dict[str, Any]:
    image = load_image(str(payload["image"]))
    corners = payload.get("corners")
    if not isinstance(corners, list) or len(corners) != 4:
        raise ValueError("Four screen corners are required.")
    grid = grid_from_payload(payload.get("grid"))
    warped = warp_screen(image, corners)
    warped = clean_warped_image(warped)  # #14/#15/#17: cursor / glare / outlines
    refined, summary = whole_grid_focused_recheck(warped, grid, str(payload.get("mode", "conservative")))
    color_grid = extract_color_semantics(warped, refined)  # #23: color labels
    return {"grid": refined, "verification": summary, "colorGrid": color_grid}


def add_table_header(table: Any) -> None:
    header = table.rows[0].cells
    header[0].text = "Row"
    for col in range(COLS):
        label = str(col) if 1 <= col <= 38 else ""
        header[col + 1].text = label
        for paragraph in header[col + 1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_docx(payload: dict[str, Any]) -> dict[str, str]:
    grid = payload.get("grid")
    if not isinstance(grid, list) or len(grid) != ROWS:
        raise ValueError("A 13-row grid is required.")

    document = Document()
    document.add_heading("777-9 MCDU Grid Extraction", level=1)
    document.add_paragraph("Rows are numbered 1 to 13. Screen columns 2 to 39 are labelled 1 to 38.")

    table = document.add_table(rows=ROWS + 1, cols=COLS + 1)
    table.style = "Table Grid"
    add_table_header(table)

    for row_index in range(ROWS):
        cells = table.rows[row_index + 1].cells
        cells[0].text = str(row_index + 1)
        for col_index in range(COLS):
            value = ""
            if col_index in (0, COLS - 1):
                value = ""
            elif row_index < len(grid) and isinstance(grid[row_index], list) and col_index < len(grid[row_index]):
                value = str(grid[row_index][col_index])
            cells[col_index + 1].text = value.strip()
            cells[col_index + 1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[col_index + 1].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(7)

    filename = f"mcdu-grid-{uuid.uuid4().hex[:8]}.docx"
    path = EXPORTS / filename
    document.save(path)
    return {"url": f"/data/exports/{filename}", "filename": filename}


def remember(payload: dict[str, Any]) -> dict[str, Any]:
    original_raw = str(payload.get("original", ""))
    corrected_raw = str(payload.get("corrected", ""))
    row = int(payload.get("row", 0))
    if not 0 <= row < ROWS:
        raise ValueError("Correction row must be between 1 and 13.")

    original = fixed_row_text(original_raw)
    corrected_list = list(fixed_row_text(corrected_raw))
    corrected_list[0] = " "
    corrected_list[COLS - 1] = " "
    corrected = "".join(corrected_list)
    if original == corrected:
        return {"count": len(load_corrections()), "skipped": True}

    corrections = load_corrections()
    corrections[correction_key(row, original)] = corrected
    save_corrections(corrections)
    return {"count": len(corrections)}


def remember_grid(payload: dict[str, Any]) -> dict[str, Any]:
    source_grid = payload.get("sourceGrid")
    corrected_grid = payload.get("grid")
    if not isinstance(source_grid, list) or not isinstance(corrected_grid, list):
        raise ValueError("Source and corrected 13-row grids are required.")

    corrections = load_corrections()
    corrections = {key: value for key, value in corrections.items() if not key.startswith("image:")}
    saved = 0
    for row in range(ROWS):
        original = grid_row_from_payload(source_grid, row)
        corrected = grid_row_from_payload(corrected_grid, row)
        if original == corrected:
            continue
        corrections[correction_key(row, original)] = corrected
        saved += 1
    save_corrections(corrections)
    return {"count": len(corrections), "saved": saved}


def grid_row_from_payload(grid: list[Any], row: int) -> str:
    values = grid[row] if row < len(grid) and isinstance(grid[row], list) else []
    cells = [str(values[col])[:1] if col < len(values) else "" for col in range(COLS)]
    cells[0] = ""
    cells[COLS - 1] = ""
    return fixed_row_text("".join(cell or " " for cell in cells))


def fuse_photos(payload: dict[str, Any]) -> dict[str, Any]:
    """API wrapper for /api/fuse-grids (#24).

    Expects ``{"grids": [grid1, grid2, ...]}`` where each grid is a 13×40
    list of lists.  Accepts 2–3 grids.
    """
    raw = payload.get("grids")
    if not isinstance(raw, list) or not (2 <= len(raw) <= 3):
        raise ValueError("Provide 2 or 3 grids in the 'grids' list.")
    for i, g in enumerate(raw):
        if not isinstance(g, list) or len(g) != ROWS:
            raise ValueError(f"Grid {i} must have {ROWS} rows.")
        for j, row in enumerate(g):
            if not isinstance(row, list) or len(row) != COLS:
                raise ValueError(f"Grid {i}, row {j} must have {COLS} columns.")
    fused, summary = fuse_grids(raw)
    return {"grid": fused, "fusion": summary}


class McmduHandler(SimpleHTTPRequestHandler):
    def end_headers(self) -> None:
        self.send_header("Cache-Control", "no-store, no-cache, must-revalidate")
        self.send_header("Pragma", "no-cache")
        self.send_header("Expires", "0")
        super().end_headers()

    def translate_path(self, path: str) -> str:
        parsed = urlparse(path).path
        if parsed.startswith("/data/exports/"):
            return str((EXPORTS / Path(parsed).name).resolve())
        if parsed == "/":
            return str(STATIC / "index.html")
        filename = Path(parsed).name
        if filename not in {"index.html", "app.js", "styles.css"}:
            return str((STATIC / "__not_found__").resolve())
        return str((STATIC / filename).resolve())

    def do_POST(self) -> None:
        try:
            payload = read_json_body(self)
            if self.path == "/api/analyze":
                json_response(self, HTTPStatus.OK, analyze(payload))
            elif self.path == "/api/detect-display":
                json_response(self, HTTPStatus.OK, detect_display(payload))
            elif self.path == "/api/flatten-display":
                json_response(self, HTTPStatus.OK, flatten_display(payload))
            elif self.path == "/api/export-docx":
                json_response(self, HTTPStatus.OK, export_docx(payload))
            elif self.path == "/api/remember":
                json_response(self, HTTPStatus.OK, remember(payload))
            elif self.path == "/api/remember-grid":
                json_response(self, HTTPStatus.OK, remember_grid(payload))
            elif self.path == "/api/remember-templates":
                json_response(self, HTTPStatus.OK, remember_templates(payload))
            elif self.path == "/api/review-requirements":
                json_response(self, HTTPStatus.OK, review_requirements(payload))
            elif self.path == "/api/refine-grid":
                json_response(self, HTTPStatus.OK, refine_grid(payload))
            elif self.path == "/api/fuse-grids":
                json_response(self, HTTPStatus.OK, fuse_photos(payload))
            else:
                json_response(self, HTTPStatus.NOT_FOUND, {"error": "Unknown endpoint."})
        except Exception as exc:  # noqa: BLE001 - API boundary should return useful errors.
            print(f"{self.path} failed: {exc}")
            json_response(self, HTTPStatus.BAD_REQUEST, {"error": str(exc)})


def main() -> None:
    ensure_dirs()
    port = int(os.environ.get("PORT", "8766"))
    server = ThreadingHTTPServer(("127.0.0.1", port), McmduHandler)
    print(f"777-9 MCDU Grid Tool running at http://127.0.0.1:{port}")
    server.serve_forever()


if __name__ == "__main__":
    main()
